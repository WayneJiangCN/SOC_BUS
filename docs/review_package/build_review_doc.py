from __future__ import annotations

import html
import math
import os
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SVG_DIR = ROOT / "editable_diagrams"
PNG_DIR = ROOT / "rendered_diagrams"
QA_DIR = ROOT / "qa_render"
DOCX_PATH = ROOT / "多核SoC总线模型方案与架构设计文档.docx"

EDGE = Path(r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe")

NAVY = "1F4E79"
BLUE = "2E74B5"
LIGHT_BLUE = "DCE6F1"
PALE_BLUE = "EEF4FA"
GRAY = "666666"
MID_GRAY = "A6A6A6"
LIGHT_GRAY = "F2F4F7"
GRID = "C9D2DC"
BLACK = "1F1F1F"
WHITE = "FFFFFF"
RED = "C00000"
AMBER = "B45F06"
GREEN = "0B6E4F"


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, twips):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(twips))
    tcW.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, total=9360, indent=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(total))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), str(indent))
    tblInd.set(qn("w:type"), "dxa")
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        cant = OxmlElement("w:cantSplit")
        trPr.append(cant)
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(i, len(widths) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def set_run_font(run, latin="Calibri", east_asia="等线", size=None, bold=None,
                 italic=None, color=None):
    run.font.name = latin
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), latin)
    rFonts.set(qn("w:hAnsi"), latin)
    rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_border_bottom(paragraph, color=BLUE, size=12, space=3):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)


def set_paragraph_shading(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = pPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        pPr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text_node = OxmlElement("w:t")
    text_node.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text_node, end])
    set_run_font(run, size=8.5, color=GRAY)


def add_toc_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "目录将在打开文档时更新。"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, placeholder, end])
    set_run_font(run, size=10, color=GRAY)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.10
    pf.widow_control = True

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, NAVY, 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.keep_together = True
        st.paragraph_format.widow_control = True

    cap = styles["Caption"]
    cap.font.name = "Calibri"
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    cap.font.size = Pt(9)
    cap.font.italic = False
    cap.font.color.rgb = RGBColor.from_string(GRAY)
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after = Pt(9)
    cap.paragraph_format.keep_with_next = False
    cap.paragraph_format.keep_together = True

    for style_name in ["List Bullet", "List Bullet 2", "List Number"]:
        st = styles[style_name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
        st.font.size = Pt(10.5)
        st.paragraph_format.left_indent = Inches(0.5)
        st.paragraph_format.first_line_indent = Inches(-0.25)
        st.paragraph_format.space_after = Pt(5)
        st.paragraph_format.line_spacing = 1.1


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def set_headers_footers(section, cover=False):
    header = section.header
    p = header.paragraphs[0]
    for child in list(p._p):
        if child.tag != qn("w:pPr"):
            p._p.remove(child)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    if cover:
        r = p.add_run("架构评审输入")
    else:
        r = p.add_run("多核SoC总线模型方案与架构设计")
    set_run_font(r, size=8.5, bold=True, color=GRAY)
    r2 = p.add_run("    |    BUS-MDL-ARCH-001")
    set_run_font(r2, size=8.5, color=MID_GRAY)

    footer = section.footer
    fp = footer.paragraphs[0]
    for child in list(fp._p):
        if child.tag != qn("w:pPr"):
            fp._p.remove(child)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    r = fp.add_run("内部技术资料  |  V0.9  |  第 ")
    set_run_font(r, size=8.5, color=GRAY)
    add_page_field(fp)
    r = fp.add_run(" 页")
    set_run_font(r, size=8.5, color=GRAY)


def add_body(doc, text, bold_prefix=None, keep=False, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_together = keep
    p.paragraph_format.widow_control = True
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True, color=BLACK)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, color=BLACK)
    else:
        r = p.add_run(text)
        set_run_font(r, color=BLACK)
    return p


def add_label_para(doc, label, text, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.left_indent = Inches(0.08)
    r = p.add_run(label)
    set_run_font(r, bold=True, color=NAVY)
    r = p.add_run(text)
    set_run_font(r, color=BLACK)
    return p


def add_bullets(doc, items, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    for item in items:
        p = doc.add_paragraph(style=style)
        r = p.add_run(item)
        set_run_font(r, size=10.5)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_table(doc, headers, rows, widths=None, font_size=8.8, header_fill=LIGHT_GRAY,
              first_col_bold=False, repeat_header=True):
    cols = len(headers)
    if widths is None:
        base = 9360 // cols
        widths = [base] * cols
        widths[-1] += 9360 - sum(widths)
    table = doc.add_table(rows=1, cols=cols)
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    hdr = table.rows[0]
    if repeat_header:
        set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        shade_cell(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(str(h))
        set_run_font(r, size=font_size, bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(str(val))
            set_run_font(r, size=font_size, bold=(first_col_bold and i == 0), color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_callout(doc, title, text, fill=PALE_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title + "\n")
    set_run_font(r, size=10.5, bold=True, color=accent)
    r = p.add_run(text)
    set_run_font(r, size=10, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_mechanism(doc, title, trigger, resources, state, blocked, release, edges, impact):
    add_heading(doc, title, 3)
    add_label_para(doc, "触发条件：", trigger)
    add_label_para(doc, "占用资源：", resources)
    add_label_para(doc, "状态变化：", state)
    add_label_para(doc, "阻塞条件：", blocked)
    add_label_para(doc, "资源释放条件：", release)
    add_label_para(doc, "边界场景：", edges)
    add_label_para(doc, "性能与精度影响：", impact, after=7)


def _esc(text):
    return html.escape(str(text), quote=True)


class SvgCanvas:
    def __init__(self, title, subtitle=""):
        self.parts = [
            '<?xml version="1.0" encoding="UTF-8"?>',
            '<svg xmlns="http://www.w3.org/2000/svg" width="1400" height="800" viewBox="0 0 1400 800">',
            '<defs>',
            '<marker id="a-blue" markerWidth="10" markerHeight="10" refX="9" refY="3" orient="auto" markerUnits="strokeWidth"><path d="M0,0 L0,6 L9,3 z" fill="#2563EB"/></marker>',
            '<marker id="a-orange" markerWidth="10" markerHeight="10" refX="9" refY="3" orient="auto" markerUnits="strokeWidth"><path d="M0,0 L0,6 L9,3 z" fill="#D97706"/></marker>',
            '<marker id="a-green" markerWidth="10" markerHeight="10" refX="9" refY="3" orient="auto" markerUnits="strokeWidth"><path d="M0,0 L0,6 L9,3 z" fill="#059669"/></marker>',
            '<marker id="a-gray" markerWidth="10" markerHeight="10" refX="9" refY="3" orient="auto" markerUnits="strokeWidth"><path d="M0,0 L0,6 L9,3 z" fill="#64748B"/></marker>',
            '<marker id="a-red" markerWidth="10" markerHeight="10" refX="9" refY="3" orient="auto" markerUnits="strokeWidth"><path d="M0,0 L0,6 L9,3 z" fill="#C2413B"/></marker>',
            '<style>.t{font-family:"Microsoft YaHei","Noto Sans CJK SC",Arial,sans-serif;fill:#1F2937}.s{font-size:18px}.m{font-size:21px;font-weight:600}.h{font-size:30px;font-weight:700}.sub{font-size:17px;fill:#475569}.tiny{font-size:15px;fill:#475569}</style>',
            '</defs>',
            '<rect x="0" y="0" width="1400" height="800" fill="#FFFFFF"/>',
            f'<text x="48" y="48" class="t h">{_esc(title)}</text>',
        ]
        if subtitle:
            self.parts.append(f'<text x="48" y="78" class="t sub">{_esc(subtitle)}</text>')

    def box(self, x, y, w, h, title, lines=None, fill="#F8FAFC", stroke="#64748B",
            title_fill=None, dash=None, radius=4, fontsize=18):
        dash_attr = f' stroke-dasharray="{dash}"' if dash else ""
        self.parts.append(f'<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="{radius}" fill="{fill}" stroke="{stroke}" stroke-width="2"{dash_attr}/>')
        if title_fill:
            self.parts.append(f'<rect x="{x}" y="{y}" width="{w}" height="34" rx="{radius}" fill="{title_fill}"/>')
            ty = y + 24
        else:
            ty = y + 29
        self.parts.append(f'<text x="{x+w/2}" y="{ty}" text-anchor="middle" class="t m" style="font-size:{fontsize}px">{_esc(title)}</text>')
        if lines:
            start = y + (54 if title_fill else 58)
            for i, line in enumerate(lines):
                self.parts.append(f'<text x="{x+w/2}" y="{start+i*24}" text-anchor="middle" class="t tiny">{_esc(line)}</text>')

    def text(self, x, y, text, cls="s", anchor="start", color=None, weight=None):
        style = ""
        if color:
            style += f"fill:{color};"
        if weight:
            style += f"font-weight:{weight};"
        self.parts.append(f'<text x="{x}" y="{y}" text-anchor="{anchor}" class="t {cls}" style="{style}">{_esc(text)}</text>')

    def line(self, points, kind="req", width=3, arrow=True, dash=None):
        styles = {
            "req": ("#2563EB", None, "a-blue"),
            "wdata": ("#D97706", "12 6", "a-orange"),
            "rsp": ("#059669", "4 5", "a-green"),
            "ctrl": ("#64748B", "12 5 2 5", "a-gray"),
            "bp": ("#C2413B", "2 6", "a-red"),
        }
        color, default_dash, marker = styles[kind]
        d = default_dash if dash is None else dash
        dash_attr = f' stroke-dasharray="{d}"' if d else ""
        arrow_attr = f' marker-end="url(#{marker})"' if arrow else ""
        pts = " ".join(f"{x},{y}" for x, y in points)
        self.parts.append(f'<polyline points="{pts}" fill="none" stroke="{color}" stroke-width="{width}" stroke-linejoin="round" stroke-linecap="round"{dash_attr}{arrow_attr}/>')

    def separator(self, x1, y1, x2, y2, color="#CBD5E1", dash=""):
        dash_attr = f' stroke-dasharray="{dash}"' if dash else ""
        self.parts.append(f'<line x1="{x1}" y1="{y1}" x2="{x2}" y2="{y2}" stroke="{color}" stroke-width="1.5"{dash_attr}/>')

    def legend(self, y=752, include_bp=True):
        x = 52
        entries = [("req", "请求"), ("wdata", "写数据"), ("rsp", "响应"), ("ctrl", "配置/统计")]
        if include_bp:
            entries.append(("bp", "反压") )
        for kind, label in entries:
            self.line([(x, y), (x+55, y)], kind=kind, width=3)
            self.text(x+68, y+6, label, cls="tiny")
            x += 190

    def finish(self):
        self.parts.append('</svg>')
        return "\n".join(self.parts)


def diagram_1():
    c = SvgCanvas("模型边界与系统上下文", "评审问题：哪些行为属于总线模型，哪些由流量源与Target服务模型提供")
    c.box(45, 135, 225, 390, "上游环境（边界外）", ["多Core/多Master", "Trace/合成流量源", "一致性/顺序适配", "完成响应接收"], fill="#F8FAFC")
    c.box(315, 105, 750, 500, "多核SoC总线性能模型（边界内）", [], fill="#F7FBFF", stroke="#2E74B5", title_fill="#DCE6F1")
    c.box(355, 175, 180, 135, "流量源适配层", ["请求规范化", "MasterID/TxnID", "OSD准入"], fill="#FFFFFF")
    c.box(575, 175, 205, 135, "互连传输平面", ["Ring路由/仲裁", "FIFO/链路带宽", "请求/数据/响应"], fill="#FFFFFF")
    c.box(820, 175, 205, 135, "Target接入层", ["Credit准入", "请求整形", "响应回注"], fill="#FFFFFF")
    c.box(355, 380, 205, 135, "地址映射与交织", ["地址域校验", "Target选择", "路径确定"], fill="#FFFFFF")
    c.box(590, 380, 205, 135, "配置与约束管理", ["参数生效检查", "单位/范围校验", "配置指纹"], fill="#FFFFFF")
    c.box(825, 380, 200, 135, "统计与追踪管理", ["延迟/吞吐", "占用率/阻塞原因", "事务路径追踪"], fill="#FFFFFF")
    c.box(1110, 135, 245, 390, "下游环境（边界外）", ["L2服务模型", "DDR服务模型", "MMIO/其他Target", "可选RTL/硅数据基准"], fill="#F8FAFC")
    c.line([(270, 225), (355, 225)], "req")
    c.line([(270, 270), (355, 270)], "wdata")
    c.line([(355, 325), (285, 325), (285, 335), (270, 335)], "rsp")
    c.line([(535, 240), (575, 240)], "req")
    c.line([(780, 240), (820, 240)], "req")
    c.line([(1025, 240), (1110, 240)], "req")
    c.line([(1025, 280), (1110, 280)], "wdata")
    c.line([(1110, 335), (1025, 335), (1025, 290), (820, 290)], "rsp")
    c.line([(820, 330), (780, 330), (780, 290), (575, 290)], "rsp")
    c.line([(575, 330), (535, 330), (535, 290), (355, 290)], "rsp")
    c.line([(457, 380), (457, 310)], "ctrl")
    c.line([(692, 380), (692, 310)], "ctrl")
    c.line([(925, 380), (925, 310)], "ctrl")
    c.line([(1110, 450), (1025, 450)], "ctrl")
    c.legend()
    return c.finish()


def diagram_2():
    c = SvgCanvas("多核双向Ring总体架构", "基线：端点接入与拓扑解耦；请求、写数据、响应为独立逻辑流量类，共享关系由链路资源配置决定")
    nodes = [(330,160),(620,120),(915,160),(1015,375),(915,575),(620,620),(330,575),(230,375)]
    labels = ["Ring Router 0","Ring Router 1","Ring Router 2","Ring Router 3","Ring Router 4","Ring Router 5","Ring Router 6","Ring Router 7"]
    for (x,y),lab in zip(nodes,labels):
        c.box(x,y,165,76,lab,["LOCAL / CW / CCW"],fill="#FFFFFF",stroke="#475569",fontsize=17)
    cw_edges = [
        ((495,198),(620,158)), ((785,158),(915,198)),
        ((1080,198),(1097,375)), ((1097,451),(1080,613)),
        ((915,613),(785,658)), ((620,658),(495,613)),
        ((330,613),(312,451)), ((312,375),(330,198)),
    ]
    for a,b in cw_edges:
        c.line([a,b],"req",width=2.5)
    ccw_edges = [
        ((620,166),(495,206)), ((915,206),(785,166)),
        ((1088,375),(1072,206)), ((1072,605),(1088,451)),
        ((785,650),(915,605)), ((495,605),(620,650)),
        ((320,451),(338,605)), ((338,206),(320,375)),
    ]
    for a,b in ccw_edges:
        c.line([a,b],"rsp",width=2.5)
    c.box(55,130,185,110,"Master网络接口 M0",["注入FIFO / OSD"],fill="#EEF4FA",stroke="#2E74B5")
    c.box(55,280,185,110,"Master网络接口 M1",["注入FIFO / OSD"],fill="#EEF4FA",stroke="#2E74B5")
    c.box(55,490,185,110,"Master网络接口 M2..N",["注入FIFO / OSD"],fill="#EEF4FA",stroke="#2E74B5")
    c.line([(240,185),(330,198)],"req")
    c.line([(240,335),(230,413)],"req")
    c.line([(240,545),(330,613)],"req")
    c.box(1160,130,190,110,"Target网络接口 T0",["Credit / 请求FIFO"],fill="#EEF8F4",stroke="#0B6E4F")
    c.box(1160,295,190,110,"Target网络接口 T1",["Credit / 请求FIFO"],fill="#EEF8F4",stroke="#0B6E4F")
    c.box(1160,490,190,110,"Target网络接口 T2..M",["Credit / 请求FIFO"],fill="#EEF8F4",stroke="#0B6E4F")
    c.line([(1080,198),(1160,185)],"req")
    c.line([(1180,350),(1098,413)],"req")
    c.line([(1080,613),(1160,545)],"req")
    c.box(500,325,400,150,"全局策略与观测（不承载旁路数据）",["地址映射与交织引擎  |  路由与仲裁策略", "配置与约束管理  |  统计与追踪管理", "所有数据必须经过Ring共享链路资源"],fill="#FAFAFA",stroke="#64748B",dash="7 5")
    c.line([(700,325),(700,250)],"ctrl")
    c.line([(900,400),(1015,413)],"ctrl")
    c.legend()
    return c.finish()


def diagram_3():
    c = SvgCanvas("请求、写数据与响应路径", "事务键 TxnKey={MasterID, LocalTxnID, Epoch} 在所有阶段保持稳定；Seq/BeatIndex关联多段返回")
    xs=[80,330,570,830,1080]
    titles=["流量源适配层","Master网络接口","Ring传输平面","Target网络接口","Target服务模型"]
    for x,t in zip(xs,titles):
        c.box(x,110,200,62,t,fill="#F8FAFC",stroke="#64748B",fontsize=17)
        c.separator(x+100,172,x+100,680,color="#CBD5E1",dash="6 6")
    # read
    c.text(50,225,"读事务",cls="m",color="#1F4E79")
    c.line([(180,220),(430,220),(670,220),(930,220),(1180,220)],"req")
    c.text(680,208,"READ_REQ / TxnKey / ExpectedRsp",cls="tiny",anchor="middle",color="#2563EB")
    c.line([(1180,275),(930,275),(670,275),(430,275),(180,275)],"rsp")
    c.text(680,264,"READ_RSP / TxnKey / Seq / Last",cls="tiny",anchor="middle",color="#059669")
    c.text(435,300,"Last响应被上游接受后释放Master OSD",cls="tiny",anchor="middle")
    c.separator(55,330,1345,330,color="#94A3B8")
    # write
    c.text(50,385,"写事务",cls="m",color="#1F4E79")
    c.line([(180,380),(430,380),(670,380),(930,380),(1180,380)],"req")
    c.text(680,368,"WRITE_REQ / TxnKey / DataLength",cls="tiny",anchor="middle",color="#2563EB")
    c.line([(1180,430),(930,430),(670,430),(430,430)],"rsp")
    c.text(795,418,"WRITE_GRANT / TxnKey / DataTag",cls="tiny",anchor="middle",color="#059669")
    c.line([(430,490),(670,490),(930,490),(1180,490)],"wdata")
    c.text(805,478,"WRITE_DATA / TxnKey / DataTag / BeatIndex",cls="tiny",anchor="middle",color="#D97706")
    c.line([(1180,550),(930,550),(670,550),(430,550),(180,550)],"rsp")
    c.text(680,538,"WRITE_COMPLETE / TxnKey / Status",cls="tiny",anchor="middle",color="#059669")
    c.text(680,600,"请求头与写数据可解耦传输，但Target提交前必须按TxnKey/DataTag闭环关联",cls="s",anchor="middle",weight="600")
    c.text(680,632,"错误、超时或取消均必须生成终结状态；不得静默丢弃或提前复用payload",cls="tiny",anchor="middle")
    c.legend()
    return c.finish()


def diagram_4():
    c = SvgCanvas("OSD、Credit、FIFO与反压协同", "端到端并发、接收能力、突发吸收和链路带宽分别建模；下游恢复时由持续vld/ready事件自动重试")
    c.box(55,210,210,250,"Master网络接口",["Master OSD=已接受未完成", "请求/数据注入FIFO", "事务追踪表", "vld保持直到握手"],fill="#EEF4FA",stroke="#2E74B5")
    c.box(320,210,210,250,"Ring Router",["输入FIFO", "输出仲裁", "队头保持", "按流量类调度"],fill="#F8FAFC")
    c.box(585,210,210,250,"Ring Link",["in-flight容量", "序列化占用", "传播延迟", "下游满则不出队"],fill="#F8FAFC")
    c.box(850,210,210,250,"Target网络接口",["Target Credit", "Target请求FIFO", "响应回注FIFO", "服务准入"],fill="#EEF8F4",stroke="#0B6E4F")
    c.box(1115,210,230,250,"Target服务模型",["固定/分布服务延迟", "读写服务带宽", "Bank/队列并发", "完成/响应产生"],fill="#EEF8F4",stroke="#0B6E4F")
    c.line([(265,365),(320,365),(530,365),(585,365),(795,365),(850,365),(1060,365),(1115,365)],"req")
    c.line([(265,395),(320,395),(530,395),(585,395),(795,395),(850,395),(1060,395),(1115,395)],"wdata")
    c.line([(1115,425),(1060,425),(850,425),(795,425),(585,425),(530,425),(320,425),(265,425),(55,425)],"rsp")
    c.line([(1115,495),(1100,495),(1100,505),(850,505),(850,515),(585,515),(585,525),(320,525),(320,535),(55,535)],"bp")
    c.text(700,565,"空间 / Credit / ready恢复",cls="tiny",anchor="middle",color="#C2413B")
    c.text(700,595,"状态变化触发重评估；vld仍为1的队头自动重试",cls="s",anchor="middle",color="#C2413B",weight="600")
    c.box(450,625,500,72,"资源释放规则",["Master OSD：终结响应被上游接受；Target Credit：Target占用真正结束；FIFO：下游握手成功"],fill="#FFF7ED",stroke="#B45F06",fontsize=18)
    c.legend(y=748)
    return c.finish()


def diagram_5():
    c = SvgCanvas("端到端延迟分解与关键瓶颈", "每一类时延只在唯一资源所有者处记账；固定传播与序列化占用不重复叠加")
    blocks=[
        (55,180,145,"源端准入",["T_osd","T_src_q"]),
        (230,180,145,"注入与仲裁",["T_inj","T_arb"]),
        (405,180,145,"逐跳Router",["ΣT_router","ΣT_hop_q"]),
        (580,180,145,"逐跳Link",["ΣT_prop","ΣT_ser"]),
        (755,180,145,"Target入口",["T_tgt_q","T_credit"]),
        (930,180,145,"Target服务",["T_service","T_bank"]),
        (1105,180,190,"响应回程",["T_rsp_q","T_rsp_path"]),
    ]
    for x,y,w,t,lines in blocks:
        c.box(x,y,w,150,t,lines,fill="#F8FAFC",stroke="#64748B",fontsize=17)
    for i in range(len(blocks)-1):
        x=blocks[i][0]+blocks[i][2]; nx=blocks[i+1][0]
        c.line([(x,255),(nx,255)],"req",width=2.5)
    c.box(85,410,270,150,"容量瓶颈：源端",["OSD < 带宽时延积", "注入端口带宽不足", "流量源无法持续供给"],fill="#FFF7ED",stroke="#B45F06")
    c.box(395,410,270,150,"容量瓶颈：Ring最小割",["热点链路利用率≈100%", "仲裁等待与FIFO占用上升", "方向/映射不均衡"],fill="#FFF7ED",stroke="#B45F06")
    c.box(705,410,270,150,"容量瓶颈：Target",["Credit长期接近0", "服务队列增长", "目标带宽/Bank冲突"],fill="#FFF7ED",stroke="#B45F06")
    c.box(1015,410,300,150,"容量瓶颈：响应路径",["读数据体量主导", "写授权/完成受阻", "回程共享资源饱和"],fill="#FFF7ED",stroke="#B45F06")
    c.line([(220,410),(220,330)],"ctrl")
    c.line([(530,410),(650,330)],"ctrl")
    c.line([(840,410),(840,330)],"ctrl")
    c.line([(1165,410),(1200,330)],"ctrl")
    c.text(700,625,"T_total = T_fixed + T_topology + T_serialization + T_arbitration + T_queue + T_target + T_return",cls="m",anchor="middle",color="#1F4E79")
    c.text(700,660,"吞吐上界 = min(源端注入，Ring最小割，Target服务，响应带宽，OSD×有效载荷/平均RTT)",cls="s",anchor="middle")
    c.legend()
    return c.finish()


def diagram_6():
    c = SvgCanvas("性能精度校准与验收闭环", "80%仅为目标；校准集与留出验收集严格隔离，结果按KPI分层而非单一平均误差解释")
    c.box(55,145,220,135,"基准定义",["硬件/RTL配置快照", "工作负载与测量窗口", "参考统计与置信区间"],fill="#F8FAFC")
    c.box(330,145,220,135,"场景与数据集",["零负载/饱和/热点", "读写混合/OSD扫描", "校准集与留出集"],fill="#F8FAFC")
    c.box(605,145,220,135,"模型执行",["配置指纹", "确定性随机种子", "原始统计与事务追踪"],fill="#EEF4FA",stroke="#2E74B5")
    c.box(880,145,220,135,"KPI对齐",["吞吐/延迟分位数", "饱和拐点/扩展趋势", "瓶颈分类/公平性"],fill="#F8FAFC")
    c.box(1155,145,190,135,"误差分析",["结构误差", "参数误差", "随机误差"],fill="#FFF7ED",stroke="#B45F06")
    for a,b in [(275,330),(550,605),(825,880),(1100,1155)]:
        c.line([(a,212),(b,212)],"ctrl")
    c.box(1000,395,270,135,"参数校准（仅校准集）",["固定时延/带宽", "队列/OSD/Credit", "Target服务分布"],fill="#FFF7ED",stroke="#B45F06")
    c.box(635,395,270,135,"模型版本冻结",["参数来源与变更记录", "敏感度检查", "禁止对留出集调参"],fill="#EEF4FA",stroke="#2E74B5")
    c.box(270,395,270,135,"留出集验收",["加权准确度≥0.80", "硬门限全部通过", "异常场景有归因"],fill="#EEF8F4",stroke="#0B6E4F")
    c.box(55,600,410,90,"评审签核 / 不通过回退",["签核：固化模型适用域与版本", "不通过：回到结构假设或基准定义"],fill="#F8FAFC")
    c.line([(1250,280),(1250,395)],"ctrl")
    c.line([(1000,462),(905,462)],"ctrl")
    c.line([(635,462),(540,462)],"ctrl")
    c.line([(405,530),(405,600)],"ctrl")
    c.line([(270,440),(180,440),(180,280),(165,280)],"bp")
    c.line([(1115,395),(1115,330),(715,330),(715,280)],"bp")
    c.text(705,346,"仅当误差被证明来自可校准参数时回调参数",cls="tiny",anchor="middle",color="#C2413B")
    c.legend()
    return c.finish()


def write_diagrams():
    SVG_DIR.mkdir(parents=True, exist_ok=True)
    PNG_DIR.mkdir(parents=True, exist_ok=True)
    diagrams = [diagram_1(), diagram_2(), diagram_3(), diagram_4(), diagram_5(), diagram_6()]
    names = [
        "01_model_boundary_context",
        "02_bidirectional_ring_architecture",
        "03_transaction_paths",
        "04_osd_credit_backpressure",
        "05_latency_bottlenecks",
        "06_calibration_loop",
    ]
    for name, svg in zip(names, diagrams):
        svg_path = SVG_DIR / f"{name}.svg"
        png_path = PNG_DIR / f"{name}.png"
        svg_path.write_text(svg, encoding="utf-8")
        if not EDGE.exists():
            raise RuntimeError(f"Microsoft Edge not found: {EDGE}")
        uri = svg_path.resolve().as_uri()
        cmd = [
            str(EDGE), "--headless", "--disable-gpu", "--hide-scrollbars",
            "--force-device-scale-factor=1", "--window-size=1400,800",
            f"--screenshot={png_path.resolve()}", uri,
        ]
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=30)
        if not png_path.exists() or png_path.stat().st_size < 10_000:
            raise RuntimeError(f"Diagram render failed: {png_path}")
    manifest = (
        "多核SoC总线模型方案与架构设计文档——可编辑图源\n"
        "格式：SVG 1.1；画布：1400×800；背景：白色；字体：Microsoft YaHei/Arial。\n"
        "路径语义：蓝色实线=请求；橙色长虚线=写数据；绿色短虚线=响应；灰色复合虚线=配置/统计；红色点线=反压。\n"
        "六张SVG均可在draw.io、Inkscape、Adobe Illustrator或文本编辑器中继续编辑。\n"
    )
    (SVG_DIR / "图源说明.txt").write_text(manifest, encoding="utf-8")
    return names


def add_figure(doc, number, name, design_purpose, modules, arrows, colors, review_question):
    add_heading(doc, f"图纸设计说明：图{number}", 3)
    add_label_para(doc, "目的：", design_purpose)
    add_label_para(doc, "模块与端口：", modules)
    add_label_para(doc, "箭头关系：", arrows)
    add_label_para(doc, "颜色与线型：", colors)
    add_label_para(doc, "专家确认问题：", review_question, after=6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(PNG_DIR / f"{name}.png"), width=Inches(6.42))
    cap = doc.add_paragraph(style="Caption")
    cap.add_run(f"图{number}  ").bold = True
    cap.add_run({
        1: "模型边界与系统上下文",
        2: "多核双向Ring总体架构",
        3: "请求、写数据与响应路径",
        4: "OSD、Credit、FIFO与反压协同",
        5: "端到端延迟分解与关键瓶颈",
        6: "性能精度校准与验收闭环",
    }[number])
    cap.add_run(f"（可编辑源：editable_diagrams/{name}.svg）")


def add_cover(doc):
    section = doc.sections[0]
    set_headers_footers(section, cover=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("ARCHITECTURE REVIEW SPECIFICATION")
    set_run_font(r, size=10, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("多核SoC总线模型方案与架构设计文档")
    set_run_font(r, size=25, bold=True, color=BLACK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("设计编码前架构评审稿")
    set_run_font(r, size=14, color=GRAY)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(14)
    set_paragraph_border_bottom(rule, color=BLUE, size=18)

    meta = [
        ("文档编号", "BUS-MDL-ARCH-001"),
        ("版本 / 状态", "V0.9 / 专家评审输入"),
        ("评审对象", "SoC架构、片上互连、存储子系统及ESL性能建模专家"),
        ("设计阶段", "模型方案与总体架构；详细设计和编码前"),
        ("基线日期", "2026-07-20"),
        ("保密属性", "内部技术资料"),
    ]
    add_table(doc, ["元数据", "内容"], meta, widths=[1800,7560], font_size=9.5, first_col_bold=True)

    add_callout(doc, "文档结论", "本稿建议采用可配置的事务级双向Ring作为首版性能模型基线，并保持端点接口、Target服务模型和拓扑实现解耦。当前结论为“有条件进入详细设计”：物理通道共享关系、顺序语义、写数据协议、Credit释放点和80%精度基准集必须在架构评审中形成书面决议，Critical事项关闭后方可全面编码。")

    add_heading(doc, "版本记录", 2)
    add_table(doc, ["版本", "日期", "状态", "说明"], [
        ("V0.9", "2026-07-20", "专家评审输入", "完成模型方案、机制规格、参数/统计、精度校准、六张架构图和独立反对者审查。"),
    ], widths=[1100,1500,1800,4960], font_size=8.8)
    doc.add_page_break()


def add_executive_summary(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("一页专家评审摘要")
    set_run_font(r, size=18, bold=True, color=BLUE)
    set_paragraph_border_bottom(p, color=BLUE, size=12)

    add_body(doc, "结论：本方案以可独立运行的多Master、多Target事务级双向Ring为基线，保留决定系统性能的一阶结构——逐跳拓扑、输出端竞争、有限FIFO、链路序列化、Master OSD、Target Credit、Target服务能力以及端到端反压；不在首版引入flit级VC分配、RTL逐拍信号或一致性协议。该抽象能够用于瓶颈归因、配置敏感度和拓扑比较，但其绝对数值必须经RTL或硅数据校准后才可签核。")
    add_body(doc, "架构主张：请求、写数据和响应采用可区分的逻辑流量类，事务使用全局唯一TxnKey闭环；所有端点数据必须经过Ring共享资源，不设置绕过路径。FIFO或接口保持vld时，若下游满则包留在原位，空间恢复所产生的状态事件重新触发仲裁与发送，无需独立retry协议。Master OSD、Target Credit、FIFO容量和链路带宽各自承担不同约束，释放点不得合并或推断。")

    add_table(doc, ["评审维度", "当前建议", "签核门槛"], [
        ("拓扑", "双向Ring作为首版基线；端点接口保持拓扑无关", "确认端点规模、布点、最短路及等距规则"),
        ("资源", "三类逻辑流量分队列；链路共享关系参数化", "确认物理等效是共享链路还是独立子网"),
        ("事务", "TxnKey+Seq/Last关联读写各阶段", "确认写请求/数据先后、顺序域和异常终结"),
        ("流控", "OSD、Credit、FIFO、序列化状态分层记账", "确认Target Credit真实释放点"),
        ("精度", "80%为多KPI加权目标，不宣称已达到", "确定参考平台、校准集、留出集与硬门限"),
    ], widths=[1300,3900,4160], font_size=8.2, first_col_bold=True)

    add_callout(doc, "建议评审结论：有条件通过", "允许在评审决议写回后进入详细设计；在Critical事项关闭前，不建议冻结数据通道微架构或开展以性能结果为依据的全面编码。可先行开展配置Schema、统计框架、事务不变量和模块接口的非争议部分。", fill="FFF7ED", accent=AMBER)

    add_body(doc, "需专家拍板的P0事项：①请求/写数据/响应的物理带宽共享模型；②顺序域、原子/屏障及错误终结范围；③写请求与写数据的许可关系；④多响应事务的Target Credit与Master OSD释放点；⑤仲裁公平性和QoS最低要求；⑥80%精度的参考平台、场景权重与验收硬门限。", keep=True, after=0)
    doc.add_page_break()


def add_toc(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("目录")
    set_run_font(r, size=18, bold=True, color=BLUE)
    set_paragraph_border_bottom(p, color=BLUE, size=10)
    add_table(doc, ["章节", "内容"], [
        ("1", "文档概述"), ("2", "建设背景与目标"), ("3", "设计范围、非目标与约束"),
        ("4", "需求工程化分析"), ("5", "模型边界与总体架构"), ("6", "模块职责与接口"),
        ("7", "读写事务和数据流"), ("8", "多核竞争、仲裁与交织"), ("9", "OSD、Credit、FIFO和反压"),
        ("10", "延迟、带宽、排队及瓶颈模型"), ("11", "配置和统计设计"),
        ("12", "精度定义、校准和验证方案"), ("13", "替代方案与设计权衡"),
        ("14", "风险、假设和待确认事项"), ("15", "专家评审议题与阶段结论"),
        ("附录A", "交付物索引"),
    ], widths=[1300,8060], font_size=9.2, first_col_bold=True)
    doc.add_page_break()


def add_sections_1_to_4(doc):
    add_heading(doc, "1 文档概述", 1)
    add_heading(doc, "1.1 文档目的与评审定位", 2)
    add_body(doc, "本文定义一个可独立运行的多核SoC总线性能模型的方案边界、总体架构、事务语义、共享资源模型、配置与统计接口，以及精度校准和验收方法。文档的输出用于架构评审和后续详细设计输入，不用于描述某个现有软件实现，也不以临时类、函数、数据结构或演示用例作为正式规格。")
    add_body(doc, "评审重点不是模型是否具备基本收发功能，而是模型能否真实形成多Master竞争、热点Target拥塞、逐级反压和可解释的性能拐点；能否在不重复计算延迟和带宽的前提下保持事务闭环；以及是否建立了可被对标数据证伪的精度定义。任何尚未由硬件规格、RTL或硅数据支撑的数值，在本文中均标识为建议初值、合理假设或待确认事项。")

    add_heading(doc, "1.2 方案基线与适用结论", 2)
    add_body(doc, "本稿建议采用事务级双向Ring作为首版互连传输平面。选择依据是：Ring能够以有限的模块复杂度保留端点位置、多跳路径、方向选择和链路最小割等结构性因素，同时比共享总线更能表达空间并行，比全连接Crossbar更接近中等规模多端点互连的布局约束，又比flit/VC级Mesh具有更低的建模和校准成本。该结论仅针对首版性能模型基线，不等价于宣称最终SoC物理互连必须采用Ring。")
    add_body(doc, "模型的端点接入层、Target服务模型、地址映射和统计接口保持拓扑无关。若评审确认端点规模、流量矩阵或物理实现约束使Ring不适用，可以替换互连传输平面而不改变上游事务语义和下游服务接口。首版模型以结构和瓶颈相关性为优先，明确不承诺RTL逐周期一致。")

    add_heading(doc, "1.3 规范性术语", 2)
    add_table(doc, ["术语", "本文定义"], [
        ("Master", "能够发起读写事务的Core、DMA或合成流量源；每个Master拥有独立身份、注入队列和OSD预算。"),
        ("Target", "被访问的L2、DDR控制器、片上存储或其他服务端模型；具有独立队列、Credit和服务能力。"),
        ("TxnKey", "事务全局关联键，建议由MasterID、LocalTxnID和Epoch组成，在请求、写数据及所有响应阶段保持不变。"),
        ("OSD", "Master侧端到端Outstanding深度，统计已被模型接受但尚未终结的事务。"),
        ("Credit", "Target侧可接收/可服务资源的显式容量，不等同于Master OSD或任意FIFO空位。"),
        ("Traffic Class", "请求、写数据、响应及可选QoS类别的逻辑调度分类；其是否共享物理等效带宽由配置决定。"),
        ("反压", "下游资源不足导致上游保持数据、停止出队并沿依赖路径传播的流控行为，不通过丢包或静默重试实现。"),
        ("精度", "相对参考平台在定义场景和KPI上的量化一致程度，不以主观‘趋势相似’替代。"),
    ], widths=[1600,7760], font_size=9, first_col_bold=True)

    add_heading(doc, "2 建设背景与目标", 1)
    add_heading(doc, "2.1 建设背景", 2)
    add_body(doc, "多Core系统的总线性能通常由端到端往返时间、并发深度、热点链路、Target服务能力和地址分布共同决定。若模型仅给每笔访问叠加固定延迟，多个Master不会在同一物理等效资源上排队，增加Core数量、改变interleave或缩小FIFO也无法产生可信的饱和行为。相反，若首版即复制完整NoC微架构，则模型开发和校准成本会显著上升，且在硬件关键参数尚未冻结时难以证明额外细节能够提升系统级结论的有效性。")
    add_body(doc, "本项目因此采用分层可配置的ESL性能模型：事务层保证身份、顺序和完成闭环；互连层表达拓扑、竞争、序列化和反压；Target层表达不同存储端点的有限服务能力；配置与统计层保证每个参数可追溯、每个瓶颈可观测。模型能够脱离完整SoC运行，用于架构探索、流量压力分析和后续RTL/硅数据校准。")

    add_heading(doc, "2.2 建设目标", 2)
    add_body(doc, "功能目标是支持多个独立Master和多个异构Target，以可配置地址映射完成读写闭环，并在下游阻塞时保证无丢包、无重复提交、无资源泄漏。性能目标是同时表达零负载延迟、负载相关排队、链路带宽、逐跳拓扑、OSD限制、Target吞吐和响应回程，使热点与最小割能够自然成为瓶颈。工程目标是建立稳定的配置Schema、统计Schema、事务追踪和不变量检查，避免参数存在但不生效或统计数据不能解释结果。")
    add_body(doc, "精度目标约为80%，但在完成对标前不得表述为已达到。本文将80%定义为多个关键KPI的加权准确度和硬门限组合，并规定校准集、留出验收集、版本冻结和异常归因流程。模型只有在参考平台、场景覆盖和验收门限获得评审确认并执行后，才允许宣称达到相应适用域内的精度目标。")

    add_heading(doc, "2.3 成功判据", 2)
    add_table(doc, ["判据类别", "通过条件", "主要证据"], [
        ("事务正确性", "所有已接受事务均终结；响应归属、分段计数、写请求/写数据关联正确；资源计数回零。", "不变量、终态检查、事务追踪"),
        ("竞争真实性", "两个及以上Master访问共享路径或Target时，出现可归因的仲裁等待、队列增长或吞吐分配。", "每输出端等待、链路利用率、Master吞吐"),
        ("反压闭环", "任一有限资源满时包保持原位；资源恢复触发自动重试；不存在旁路或丢包。", "FIFO占用、vld/ready、阻塞原因、恢复事件"),
        ("性能可解释", "吞吐上界、饱和拐点和p50/p95延迟变化能由唯一资源或路径解释。", "延迟分解、利用率、stall attribution"),
        ("参数有效性", "关键参数变化产生方向正确、边界明确且可重复的影响；无死参数。", "单参数扫描、配置指纹、敏感度报告"),
        ("精度验收", "留出集加权准确度达到目标，全部硬门限通过，失败场景有结构性归因。", "对标报告、误差矩阵、版本签核"),
    ], widths=[1500,4500,3360], font_size=8.6, first_col_bold=True)

    add_heading(doc, "3 设计范围、非目标与约束", 1)
    add_heading(doc, "3.1 设计范围", 2)
    add_body(doc, "模型范围从Master请求在流量源适配层被接受开始，到最终终结响应被上游接受结束。范围内包括：Master网络接口、地址映射与交织引擎、Ring Router与Link、Target网络接口、Target服务模型的性能抽象、配置与约束管理、统计与追踪管理。范围内所有影响吞吐或延迟的有限资源必须有唯一所有者、容量、申请点、释放点和观测指标。")
    add_body(doc, "模型支持单Core和多Core流量、均匀和热点地址分布、读写混合、不同事务长度、不同OSD/FIFO/Credit、不同仲裁和交织策略、持续下游反压，以及L2和DDR等具有不同服务能力的Target。模型运行不依赖完整SoC，但允许接入trace、合成流量源或更高层ESL组件。")

    add_heading(doc, "3.2 非目标", 2)
    add_body(doc, "首版不复制完整一致性协议、目录与snoop网络，不承担软件可见内存模型的最终定义，不模拟PHY训练、DDR命令级调度、模拟电气效应或功耗热行为。首版也不要求flit级VC分配、逐拍RTL总线信号、CDC亚稳态或物理拥塞布线。若这些因素经敏感度或对标证明对目标KPI贡献超过误差预算，应以增量特性进入后续版本，而不是在未证明收益前加入基线。")
    add_body(doc, "模型输出不能单独用于签核RTL微流水优化、单拍旁路或细粒度VC策略；不能将某个建议默认值解释为真实硬件参数；不能将模型内的逻辑通道数量直接等价为物理线数。功能一致性、缓存一致性和正式协议合规需要由各自验证环境承担。")

    add_heading(doc, "3.3 约束与适用域", 2)
    add_body(doc, "基线假设单一模型时钟和统一cycle时间基准，跨时钟域通过端口适配层折算为显式固定延迟、吞吐和有限队列；若必须研究异步相位效应，应升级为多时钟事件模型。流量以单播读写为主，原子、屏障、多播和一致性事务在未获得协议定义前均列为非目标或待确认。地址与长度必须被映射引擎验证，跨interleave粒度或跨Target边界的请求应在源端分段并以父事务上下文聚合完成。")
    add_body(doc, "模型采用事件驱动、周期量化的状态推进。任何队列或接口处于vld状态时，其队头在下游状态变化、链路可用时刻、Credit返还或Target完成事件发生后重新参与调度；不是每周期盲轮询，也不需要独立retry消息。该约束既保证持续反压后的进度，也避免因缺少唤醒事件导致包永久停留。")

    add_heading(doc, "4 需求工程化分析", 1)
    add_heading(doc, "4.1 已知需求", 2)
    add_table(doc, ["编号", "需求陈述", "架构响应", "验证方式"], [
        ("REQ-01", "多Core/多Master、脱离完整SoC独立运行", "流量源适配层和每Master独立上下文；Target以可替换服务模型连接。", "多源并发、隔离与回程归属检查"),
        ("REQ-02", "建模延迟、交织、多核竞争和瓶颈", "逐跳Router/Link、地址映射、输出仲裁、有限Target资源。", "零负载、饱和、热点与敏感度场景"),
        ("REQ-03", "延迟、带宽、仲裁、FIFO、OSD可配置", "配置Schema映射到唯一资源所有者；加载时执行范围和冲突校验。", "参数作用点审计和单参数扫描"),
        ("REQ-04", "模拟L2/DDR服务能力和反压", "Target服务模型具备延迟、吞吐、并发、队列和可选Bank冲突。", "Target限速与反压传播检查"),
        ("REQ-05", "优先评估Ring并与替代方案权衡", "双向Ring基线、拓扑无关端点；保留Shared Bus/Crossbar/Mesh比较。", "相同流量和端点配置对比"),
        ("REQ-06", "性能精度目标约80%", "多KPI评分、硬门限、校准/留出集隔离。", "对标与验收报告"),
        ("REQ-07", "FIFO/接口持续vld自动重试", "下游不可接收时不出队，状态变化事件重新调度；不引入独立retry协议。", "持续阻塞后恢复、无丢包/无死停"),
    ], widths=[900,2750,3500,2210], font_size=8.2, first_col_bold=True)

    add_heading(doc, "4.2 已确定决策", 2)
    add_body(doc, "已确定决策是：本阶段交付模型方案和总体架构，不以现有代码结构作为规格；模型必须能够独立产生或接收多Master流量；关键延迟和容量参数可配置；L2、DDR等Target通过服务模型表达；80%仅作为待对标目标；FIFO或接口保持vld时由事件机制自动重试，下游满时数据原地保留，不增加独立retry协议。")
    add_body(doc, "本稿进一步形成的方案决策是：首版以双向Ring作为评审基线，端点接口保持拓扑无关；事务通过TxnKey闭环关联；资源限制按Master OSD、Target Credit、FIFO容量和链路带宽分层拥有；配置和统计属于架构的一部分而非调试附属物。上述方案决策需通过本次专家评审后转化为项目基线。")

    add_heading(doc, "4.3 合理假设及影响", 2)
    add_table(doc, ["假设", "采用原因", "若不成立的影响与处理"], [
        ("单一cycle时间基准", "便于统一延迟、吞吐和事件顺序。", "若存在多个独立频域，应增加CDC队列、频率比和相位模型。"),
        ("以单播读写为首版主流量", "项目输入未给出一致性、原子和多播协议。", "若需要这些事务，必须扩展Traffic Class、顺序依赖和完成条件。"),
        ("确定性最短路径", "可复现、易校准，避免自适应路由引入额外状态。", "若硬件采用自适应路由，需要增加拥塞选择和稳定性验证。"),
        ("协议段为最小可仲裁单位", "保留头/数据体量差异，同时避免flit级复杂度。", "若flit交错对结果显著，应引入分段或VC级扩展。"),
        ("三类逻辑流量独立排队", "降低HOL和协议依赖死锁风险。", "其物理带宽是否共享仍需配置，否则可能高估吞吐。"),
        ("目标服务可由延迟、带宽、并发和可选Bank冲突逼近", "满足约80%目标所需的一阶行为。", "若DDR调度或cache miss行为主导误差，应替换为更高保真Target模型。"),
        ("错误以终结响应闭环", "避免OSD/Credit泄漏。", "若硬件存在取消、重放或超时恢复，需增加明确状态机。"),
    ], widths=[2200,3200,3960], font_size=8.4, first_col_bold=True)

    add_heading(doc, "4.4 待专家确认事项", 2)
    add_table(doc, ["优先级", "待确认问题", "保守默认", "不确认的影响"], [
        ("P0", "请求、写数据、响应是否共享同一物理等效链路带宽", "三类逻辑队列，共享每方向链路调度器和带宽预算", "决定饱和吞吐、读写互扰及死锁规避方式"),
        ("P0", "写请求与写数据的先后/许可协议", "WRITE_REQ→WRITE_GRANT→WRITE_DATA→WRITE_COMPLETE", "影响写缓冲、DataTag、OSD和Credit生命周期"),
        ("P0", "顺序域、同地址hazard、原子与屏障范围", "仅保证单队列FIFO与TxnKey归属，不提供全局顺序", "可能导致功能语义和性能依赖不一致"),
        ("P0", "Target Credit释放于服务完成、响应生成还是响应被接受", "在Target占用真正结束时释放；与Master OSD分离", "错误释放会高估并发或造成资源泄漏"),
        ("P0", "80%参考平台、KPI权重和验收场景", "本文建议的加权评分与硬门限", "无基准则精度目标不可验证"),
        ("P1", "仲裁策略、QoS和最大饥饿时间", "RR+年龄逃逸；无带宽保证", "影响公平性和尾延迟"),
        ("P1", "Ring端点布点、等距路由和方向偏置", "最短路；等距采用基于TxnKey的确定性散列", "影响热点链路和可复现性"),
        ("P1", "协议头、数据beat和ECC开销", "头部和数据分别按配置字节数计费", "直接影响序列化与有效带宽"),
    ], widths=[800,3100,3500,1960], font_size=8.0, first_col_bold=True)


def add_sections_5_to_6(doc, names):
    add_heading(doc, "5 模型边界与总体架构", 1)
    add_heading(doc, "5.1 模型边界", 2)
    add_body(doc, "总线模型边界从Master侧请求被流量源适配层接受开始，到读数据或写完成等终结响应被上游接受结束。上游Core微架构、cache一致性、软件访存语义和流量生成逻辑位于边界外；下游DRAM命令级控制、PHY和真实存储介质位于边界外。边界内负责保存事务上下文、执行地址映射、争用共享资源、模拟Target可见服务能力、生成响应并收集可解释统计。")
    add_body(doc, "Target服务模型虽然作为可替换组件连接在总线传输平面之后，但其性能参数和反压属于本项目的评估闭环。其接口必须明确请求被接受、服务开始、服务完成、响应生成和响应被接收的时点，避免将Target延迟同时计入Target网络接口和服务模型。")
    add_figure(doc, 1, names[0],
               "固定模型责任边界，防止把流量源行为、总线竞争和Target服务混为单一延迟。",
               "边界外包含多Core/多Master及L2/DDR等基准来源；边界内包含流量源适配层、互连传输平面、Target接入层、地址映射与交织引擎、配置与约束管理、统计与追踪管理。Master侧端口为Master Request/Response Interface，Target侧端口为Target Request/Response Interface。",
               "请求和写数据从左向右进入Target；响应从右向左返回；配置与统计为控制路径，不承载数据旁路。",
               "蓝色实线为请求，橙色长虚线为写数据，绿色短虚线为响应，灰色复合虚线为配置/统计，红色点线为反压。",
               "Target服务模型是否计入总线模型交付边界，以及跨时钟域适配是否纳入首版。")

    add_heading(doc, "5.2 总体架构", 2)
    add_body(doc, "总体架构由流量源适配层、Master网络接口、地址映射与交织引擎、Ring Router、Ring Link、Target网络接口、Target服务模型、配置与约束管理、统计与追踪管理构成。Master网络接口将上游事务转换为统一协议段并实施源端OSD准入；地址映射与交织引擎决定TargetID和目的节点；Ring Router执行逐跳路由与输出仲裁；Ring Link承担唯一的逐方向带宽、传播延迟和在途容量；Target网络接口执行Credit准入、排队和响应回注。")
    add_body(doc, "双向Ring由相邻节点之间的两条有向链路构成。一个节点允许挂接一个或多个逻辑端点，但端点到本地Router的注入和弹出能力必须显式建模，不能假定无限带宽。路由默认采用确定性最短路径；等距时建议对TxnKey做稳定散列以减少固定方向偏置，同时保留可重复性。多个不共享输出链路的传输可以并行，只有竞争同一输出、同一链路调度器或同一Target资源时才相互阻塞。")
    add_figure(doc, 2, names[1],
               "说明多Master如何进入Ring、双向逐跳资源如何形成空间并行与共享瓶颈，以及控制平面不构成旁路。",
               "Master网络接口通过LOCAL注入Ring Router；Router具有LOCAL、CW和CCW端口；有向Ring Link连接相邻Router；Target网络接口从LOCAL弹出并连接Target服务模型。",
               "请求、写数据和响应沿确定性方向逐跳转发；控制平面仅向映射、路由、仲裁和统计模块提供参数或采样。",
               "路径线型与图1一致；顺时针/逆时针采用相同资源语义，不以装饰线表达。",
               "端点布点、方向带宽是否对称、等距散列规则，以及三类流量是否共享链路调度器。")

    add_heading(doc, "5.3 执行与时间语义", 2)
    add_body(doc, "模型采用离散事件驱动、周期量化的执行语义。状态变化事件包括新vld到达、FIFO由满转非满、Credit返还、链路序列化窗口开放、在途段到达、Target服务完成和响应被接受。事件发生时只重新评估受影响资源，保证在高阻塞场景下仿真性能可控。所有同一cycle事件必须使用确定性的阶段顺序或两阶段提交，防止软件调用顺序无意构成仲裁优先级。")
    add_body(doc, "建议每个cycle按‘完成/资源返还→到达入队→候选收集→仲裁决策→原子提交→统计更新’推进。候选收集与提交分离，使多个输入竞争同一输出时只产生一个winner；资源是否可用以提交前快照判断；提交后统一更新FIFO、Credit和序列化状态。若采用异步事件优化，也必须保持与该逻辑顺序等价。")

    add_heading(doc, "6 模块职责与接口", 1)
    add_heading(doc, "6.1 模块职责分解", 2)
    add_table(doc, ["模块", "核心职责", "拥有的状态/资源", "禁止承担"], [
        ("流量源适配层", "规范化上游请求，提供MasterID、顺序属性和完成回调。", "上游握手、可选父事务分段上下文。", "不得注入绕过Master网络接口的流量。"),
        ("Master网络接口", "OSD准入、协议段生成、响应聚合和终结。", "请求/写数据FIFO、事务追踪表、Master OSD。", "不得直接修改Target Credit或Ring链路状态。"),
        ("地址映射与交织引擎", "地址域校验、Target选择、分段边界和节点映射。", "地址图、interleave规则、映射版本。", "不得承担排队、带宽或Target服务延迟。"),
        ("Ring Router", "输入候选收集、下一跳选择和每输出仲裁。", "端口FIFO、仲裁指针、年龄/QoS状态。", "不得保存端到端事务完成状态。"),
        ("Ring Link", "传播、序列化和有限在途容量。", "next_free_cycle、in-flight队列、每方向带宽。", "不得重复计入Router或Target固定延迟。"),
        ("Target网络接口", "Target Credit准入、请求整形、响应回注。", "请求/响应FIFO、Credit镜像、DataTag关联。", "不得替代Target内部服务模型。"),
        ("Target服务模型", "表达L2/DDR等服务延迟、吞吐、并发和可选Bank冲突。", "服务队列、busy状态、内部资源占用。", "不得隐式承担Ring传输时延。"),
        ("配置与约束管理", "加载、继承、校验并冻结配置，生成配置指纹。", "Schema版本、参数来源和生效清单。", "不得在运行中静默修正非法参数。"),
        ("统计与追踪管理", "采样KPI、阻塞原因、占用率和事务路径。", "计数器、直方图、trace缓冲。", "不得改变数据通路调度决策。"),
    ], widths=[1600,2600,2800,2360], font_size=7.8, first_col_bold=True)

    add_heading(doc, "6.2 接口与事务载荷", 2)
    add_body(doc, "Master Request Interface承载READ_REQ、WRITE_REQ和必要的顺序/QoS属性；Master Response Interface承载READ_RSP、WRITE_GRANT、WRITE_COMPLETE和ERROR_COMPLETE。Ring Request Channel、Ring Write Data Channel和Ring Response Channel是逻辑流量类，不预先假定物理独立。Target Request Interface和Target Response Interface分别表示Target网络接口与Target服务模型之间的准入和完成边界。Control/Configuration Interface只传递静态配置、运行控制和统计采样，不提供数据旁路。")
    add_table(doc, ["字段", "建议宽度/类型", "语义与约束"], [
        ("MasterID", "配置决定；整数", "全局唯一标识发起端，参与TxnKey。"),
        ("LocalTxnID", "每Master独立；整数", "在Epoch未回绕前不得与在途事务重复。"),
        ("Epoch", "可选小整数", "解决LocalTxnID回绕；若ID空间保证不回绕可省略。"),
        ("TargetID", "地址映射输出", "在请求分段后固定，响应原样返回。"),
        ("TrafficClass", "枚举", "至少区分Request、WriteData、Response；可扩展QoS。"),
        ("Address/Length", "字节地址/字节数", "跨映射边界时必须先分段。"),
        ("Seq/BeatIndex/Last", "整数/布尔", "标识多段数据和多响应完成条件。"),
        ("ExpectedRsp", "正整数", "父事务预计终结所需的响应数，发出后不可静默更改。"),
        ("OrderDomain/QoS", "可选枚举", "只有在评审定义后参与仲裁或完成提交。"),
        ("Status", "枚举", "OK、ERROR、TIMEOUT、CANCEL等终结状态；首版至少支持OK/ERROR。"),
    ], widths=[1700,2100,5560], font_size=8.5, first_col_bold=True)

    add_heading(doc, "6.3 资源所有权与接口不变量", 2)
    add_body(doc, "每类有限资源只有一个权威所有者：Master OSD由Master网络接口拥有；Router FIFO和仲裁状态由对应Ring Router拥有；链路序列化和在途容量由Ring Link拥有；Target Credit由Target网络接口与Target服务模型按协议共同维护但只有一个权威计数；Target服务并发由Target服务模型拥有。其他模块可以读取镜像或统计，但不得直接增减权威计数。")
    add_body(doc, "接口采用保留式握手语义：发送方在vld=1且ready=0期间必须保持事务身份、类型、长度和payload引用稳定；只有握手成功后才允许出队或复用payload。接收方ready必须由实际容量决定，不能先接受再依靠丢弃恢复。任一接口状态变化都需要触发与其相邻的调度器重新评估，保证持续vld能够在下游恢复后自动前进。")


def add_section_7(doc, names):
    add_heading(doc, "7 读写事务和数据流", 1)
    add_heading(doc, "7.1 事务标识、分段与关联", 2)
    add_body(doc, "每个被模型接受的父事务由TxnKey={MasterID, LocalTxnID, Epoch}唯一标识。地址映射导致的分段、写数据beat和多路读响应在TxnKey基础上使用SegmentID、BeatIndex或Seq编号。响应必须携带TxnKey及其分段信息，禁止依赖返回顺序查找事务。事务追踪表记录请求类型、地址、总长度、Target集合、ExpectedRsp、ReceivedRsp、错误状态和资源占用标志，任何响应只能使其对应事务的单一状态边发生一次。")
    add_body(doc, "跨interleave粒度、Target地址域或最大传输长度的访问由流量源适配层或Master网络接口先分割为子事务。父事务只有在全部子事务终结且响应被上游接受后完成；任一子事务错误必须按协议聚合为父事务终结状态。ExpectedRsp在事务首次接受时确定，若Target内部再拆分，必须通过显式接口更新期望计数且保证原子性，不能由最后到达的响应隐式猜测。")

    add_heading(doc, "7.2 读事务", 2)
    add_body(doc, "读事务从Master Request Interface进入。Master网络接口检查格式、地址映射和读OSD；接受后分配事务上下文并将READ_REQ置入请求FIFO。请求经过本地注入、逐跳Ring和Target网络接口，在Target Credit可用时提交Target服务模型。Target完成后生成一个或多个READ_RSP，响应沿Ring Response Channel返回。每个响应被Master网络接口接收时更新ReceivedRsp和错误状态；只有ReceivedRsp等于ExpectedRsp且所有Last/Seq一致性检查通过，并且最终聚合响应被上游接受，事务才进入COMPLETED并释放Master OSD。")
    add_mechanism(doc, "读事务准入与完成机制",
                  "上游以vld提交合法READ_REQ，且地址可映射、事务ID不冲突。",
                  "接受时占用一个读Master OSD、一个事务追踪表项及请求FIFO表项；到达Target后再占用Target Credit和服务队列。",
                  "NEW→ACCEPTED→IN_RING→AT_TARGET→RESPONDING→COMPLETED/ERROR；每个响应使ReceivedRsp单调增加。",
                  "读OSD耗尽、请求FIFO满、下游链路不可接收、Target Credit为0、响应FIFO满或上游不接收。",
                  "Target Credit在Target占用真正结束时释放；读Master OSD仅在最后一个终结响应被上游接受时释放；各级FIFO在握手成功后释放。",
                  "零长度、跨Target分段、重复/越界Seq、响应丢失、错误与正常响应混合、ID回绕和上游长期不接收必须有确定处理。",
                  "过早释放OSD会高估源端并发，过晚释放Target Credit会低估Target吞吐；多响应计数错误会导致资源泄漏或提前完成。")

    add_heading(doc, "7.3 写事务", 2)
    add_body(doc, "写事务的请求头和写数据在逻辑上解耦，但必须由TxnKey和可选DataTag绑定。为避免Target在未知缓冲能力下接收无处存放的数据，本稿建议首版采用WRITE_REQ→WRITE_GRANT→WRITE_DATA→WRITE_COMPLETE。WRITE_REQ获得Target侧写上下文后返回WRITE_GRANT及DataTag；Master网络接口随后发送一个或多个WRITE_DATA；Target只有在全部数据接收并完成服务后返回WRITE_COMPLETE。若评审确认硬件允许请求与数据同时或数据先行，则必须同时定义数据暂存容量、乱序匹配、超时和Credit占用。")
    add_body(doc, "写Master OSD在WRITE_REQ被源端接受时占用，并贯穿授权、数据和完成阶段。WRITE_GRANT不是终结响应，不释放Master OSD；Target写Credit是否在请求阶段占用、在数据接收后释放还是在服务完成后释放，必须与硬件缓冲定义一致。建议默认在Target网络接口成功建立写上下文时占用，并在Target不再持有该写事务的命令、数据或服务资源时释放。")
    add_mechanism(doc, "写请求、写数据与完成机制",
                  "上游提交WRITE_REQ；写授权返回后写数据FIFO具有对应TxnKey/DataTag的vld数据。",
                  "写Master OSD、事务追踪表、写请求FIFO；Target准入后占用写Credit/上下文；WRITE_DATA占用数据FIFO、链路字节和Target数据缓冲。",
                  "NEW→REQ_ACCEPTED→GRANTED→DATA_IN_PROGRESS→SERVICE→COMPLETED/ERROR；DataBeatCount和ExpectedBeat必须单调闭合。",
                  "写OSD或请求FIFO满、未获得Grant、DataTag不匹配、写数据FIFO/链路/Target数据缓冲满、Target服务不可用或完成响应受阻。",
                  "请求FIFO和数据FIFO各自在相应握手后释放；Target写Credit按已确认的真实占用结束点释放；写Master OSD在WRITE_COMPLETE/ERROR被上游接受后释放。",
                  "请求无数据、部分数据后超时、重复beat、Grant丢失、错误响应、跨Target写、上游撤销以及Target复位均需终结或排空规则。",
                  "Grant往返会增加小写事务延迟但使Target缓冲约束可解释；若硬件实际上不需要Grant，校准可选择并行模式，不能用缩短固定延迟掩盖协议差异。")

    add_heading(doc, "7.4 请求、写数据和响应路径", 2)
    add_figure(doc, 3, names[2],
               "明确读写事务的阶段、方向、关联字段和完成时点，审查多响应与两阶段写是否闭环。",
               "流量源适配层、Master网络接口、Ring传输平面、Target网络接口和Target服务模型依次连接；各阶段使用Master Request/Response Interface、Ring Request/Write Data/Response Channel和Target Request/Response Interface。",
               "READ_REQ和WRITE_REQ向Target传输；WRITE_GRANT先返回；WRITE_DATA再向Target传输；READ_RSP和WRITE_COMPLETE返回Master。",
               "蓝色实线、橙色长虚线、绿色短虚线分别表示请求、写数据和响应；灰色控制路径未画入主事务序列。",
               "写Grant是否为真实协议阶段、读多响应的粒度及错误终结是否需要单独响应类型。")

    add_heading(doc, "7.5 事务完成与资源不变量", 2)
    add_table(doc, ["不变量", "检查时点", "失败含义"], [
        ("Accepted = Completed + InFlight + Aborted_with_terminal_response", "周期检查与仿真结束", "事务丢失、重复终结或计数域不一致"),
        ("0 ≤ ReceivedRsp ≤ ExpectedRsp", "每个读响应到达", "重复响应、ExpectedRsp错误或上下文污染"),
        ("0 ≤ ReceivedBeat ≤ ExpectedBeat", "每个写数据beat到达", "数据重放、越界或错误关联"),
        ("每个资源acquire仅对应一次release", "资源状态边变化", "OSD/Credit/FIFO泄漏或过释放"),
        ("无TxnKey重复处于在途", "新请求接受", "ID空间/回绕管理失效"),
        ("payload在握手或终结前保持有效", "每个阻塞周期", "引用悬空、数据被复用或不确定行为"),
        ("idle时所有队列、在途链路和事务表为空", "排空/结束", "唤醒丢失、死锁或资源泄漏"),
    ], widths=[4100,2300,2960], font_size=8.5, first_col_bold=True)


def add_section_8(doc):
    add_heading(doc, "8 多核竞争、仲裁与交织", 1)
    add_heading(doc, "8.1 竞争形成与共享资源", 2)
    add_body(doc, "多核竞争必须发生在可识别的共享资源，而不能只表现为多个Master各自独立排队。共享点包括：同一Router输出端、同一方向Ring Link调度器、同一Target网络接口准入端、同一Target服务队列以及响应回程的共享链路。每个Master具有独立源端队列和OSD，但一旦路由汇聚到同一输出或Target，候选必须进入同一仲裁域。不同方向且不共享链路的事务可以并行，这也是Ring相对共享总线的空间并行来源。")
    add_body(doc, "模型禁止以下竞争绕过：为每个Master复制独立的Target带宽；以事件调用先后替代显式输出仲裁；在链路满时直接把包交给目的端；请求和响应分别使用无限带宽旁路而未在配置中声明；或在热点Target前只增加固定惩罚而不形成队列。热点瓶颈必须可在链路利用率、仲裁等待、FIFO占用、Credit耗尽或Target busy指标中被定位。")

    add_heading(doc, "8.2 路由与路径选择", 2)
    add_body(doc, "默认路由为双向Ring上的确定性最短路径。路径计算以源节点、目的节点和Ring节点数为输入，返回CW或CCW方向及跳数。等距路径若固定选择单一方向会对对称流量形成系统性偏置，建议按TxnKey稳定散列选择方向，并提供‘固定CW、固定CCW、稳定散列’三种可配置策略。路径在事务段注入后保持不变，以避免中途改道带来的重排和依赖环。")
    add_mechanism(doc, "逐跳路由与链路转发机制",
                  "Router输入FIFO队头vld，且目的节点和Traffic Class有效。",
                  "候选占用一个输入队头；仲裁成功后占用目标输出、下游FIFO表项及Link序列化窗口。",
                  "队头由WAIT_ROUTE进入WAIT_OUTPUT，winner提交后进入IN_LINK；到达下一跳后进入新的输入FIFO或本地弹出。",
                  "下游FIFO满、Link在途容量满、next_free_cycle未到、同输出存在更高优先级winner，或本地Target/NI不可接收。",
                  "输入FIFO表项在下游握手成功时释放；Link在途表项在到达并被下游接受时释放；序列化窗口按包字节数自然到期。",
                  "源等于目的、等距路径、环节点数为1/2、超大协议段、配置热更新和目的端复位必须定义。",
                  "路径选择改变热点位置和跳数；若输出仲裁不是原子winner选择，仿真事件顺序会造成不可校准的公平性误差。")

    add_heading(doc, "8.3 输出仲裁与公平性", 2)
    add_body(doc, "仲裁分为源端注入仲裁、Router输出仲裁和Target准入仲裁。默认建议每个输出使用按Traffic Class分层的轮询：先在可服务流量类之间按可配置权重选择，再在同类输入之间Round-Robin；等待年龄超过starvation_threshold时进入逃逸优先级。固定优先级只用于复现明确的硬件QoS，不能作为通用默认。仲裁只对本周期资源可提交的候选排序，不应选择下游不可接收的候选并阻塞其他可前进候选，除非硬件存在严格HOL行为且需要建模。")
    add_mechanism(doc, "输出仲裁机制",
                  "两个或以上有效候选请求同一输出，或单一候选与链路可用窗口同时到达。",
                  "输出端本周期唯一grant、仲裁指针、权重余额、等待年龄及下游容量。",
                  "候选进入ARBITRATING；winner原子提交并更新指针/权重；loser保持队头、年龄增加且vld不撤销。",
                  "输出已被占用、下游不可接收、Traffic Class被带宽门控，或严格顺序依赖禁止越过。",
                  "grant在成功握手后释放；未提交的临时选择不消耗资源；年龄在成功或事务撤销时清零。",
                  "单候选、同周期新到达、权重为0、年龄溢出、响应优先导致请求饥饿和反压环需要验证。",
                  "策略直接影响公平性、p95/p99和混合读写吞吐。应报告每Master服务份额、最大等待和Jain指数，不能只看总带宽。")

    add_heading(doc, "8.4 地址交织与热点", 2)
    add_body(doc, "地址映射先执行地址域命中，再执行interleave。对连续线性交织，设粒度G、参与Target数N、地址域起点B，则stripe=floor((Address-B)/G)，TargetIndex=stripe mod N；Target内偏移由stripe/N和块内偏移组成。G必须是最小传输对齐的整数倍，建议限制为2的幂；事务跨G边界时先分段，禁止一个协议段在两个Target之间隐式拆分。可选XOR散列用于降低规则步长流量的别名热点，但会降低地址到Target的直观可解释性。")
    add_mechanism(doc, "地址解码与交织机制",
                  "请求完成格式检查并携带地址、长度和映射版本。",
                  "不占用传输容量；使用只读地址图、interleave配置和可选父事务分段上下文。",
                  "地址被解析为TargetID、目的节点、Target局部地址和Segment列表；结果在该段生命周期内不可改变。",
                  "地址无命中、重叠地址域、粒度非法、Target离线或分段数量超出上限。",
                  "映射本身无运行资源；父事务分段上下文随全部子事务终结释放。",
                  "未对齐、跨4KB/缓存行/Target边界、N非2次幂、地址回绕、动态重映射和XOR冲突需单独验证。",
                  "交织决定Target负载和Ring路径，是形成或消除热点的主要杠杆；错误的默认Target回退会掩盖地址配置缺陷，建议非法地址直接报错。")

    add_heading(doc, "8.5 顺序、死锁与饥饿控制", 2)
    add_body(doc, "基线仅保证同一输入FIFO内的顺序和TxnKey关联，不宣称跨Traffic Class、跨Target或跨路径的全局顺序。若协议要求同一OrderDomain内有序提交，建议在Master网络接口的完成提交端维护reorder context，而不是阻止整个Ring并行传输。same-address hazard、读后写、写后读、barrier和atomic均需要协议级定义后才能进入模型。")
    add_body(doc, "为降低协议依赖死锁风险，请求、写数据和响应至少使用独立逻辑FIFO，并为终结响应提供进度优先或保留容量；即使三类流量共享物理等效链路，也应通过独立队列和仲裁避免响应被其依赖的请求永久堵塞。死锁检查基于资源等待图和全局无进展监视器：若存在在途事务、无Target长服务且连续超过deadlock_watchdog_cycles无任何握手或状态前进，则输出等待环快照并判失败。饥饿由每候选最大等待和服务间隔观测，年龄逃逸只提供有限等待保障，不等价于QoS带宽保证。")


def add_section_9(doc, names):
    add_heading(doc, "9 OSD、Credit、FIFO和反压", 1)
    add_heading(doc, "9.1 分层流控原则", 2)
    add_body(doc, "OSD、Credit、FIFO和Link序列化分别约束端到端并发、Target接收能力、短时突发吸收和长期链路吞吐。它们可以同时阻塞同一事务，但不得由一个模糊的‘outstanding’参数替代。Master OSD在源端接受事务时申请，保证源端不能无限制造在途请求；Target Credit在Target实际承诺资源时申请，反映目标端可同时保留的上下文或数据；FIFO只有握手缓存语义，不应被当作Target Credit；Link的next_free_cycle和in-flight容量共同表达流水化带宽与传播占用。")
    add_body(doc, "资源申请与释放必须配对且带事务身份。若Target服务完成后即可释放内部slot，即使响应仍在回程，Target Credit可以早于Master OSD释放；如果Target必须保留响应数据直到被Ring接受，则Credit应延后释放。两者不同步是允许的，但必须来自真实资源生命周期，而不能为提高吞吐随意选择。")
    add_figure(doc, 4, names[3],
               "审查四类流控的边界、释放点和反压恢复机制，防止资源泄漏、双重限速和唤醒丢失。",
               "Master网络接口、Ring Router、Ring Link、Target网络接口和Target服务模型依次连接；模块内部标出OSD、FIFO、Link状态和Target Credit。",
               "数据路径向下游推进；下游容量不足时红色反压向上游传播。空间、ready或Credit恢复产生状态变化事件，重新调度仍保持vld的队头。",
               "蓝/橙/绿分别为请求、写数据和响应；红色点线为反压；灰色为资源控制。",
               "Target Credit的权威所有者和释放点、响应保留容量、以及共享链路Traffic Class的最小进度保证。")

    add_heading(doc, "9.2 OSD机制", 2)
    add_mechanism(doc, "Master OSD申请、占用与释放",
                  "合法父事务到达Master网络接口并准备被模型接受。",
                  "按Master、读写类型和可选OrderDomain占用一个OSD计数及事务追踪表项。",
                  "accepted_count增加，事务进入ACCEPTED；读写子段共享父事务OSD或按协议显式选择子事务计数。",
                  "对应OSD达到上限、事务ID冲突或追踪表无表项。",
                  "全部子事务终结且最终响应被上游接受，或错误/取消通过终结响应闭环后释放。",
                  "ExpectedRsp为0、上游不接收、超时、复位排空、父子事务计数和ID回绕。",
                  "OSD小于带宽时延积会限制吞吐；以注入成功而非上游接受作为申请点会隐藏源端排队，必须在报告中固定统计口径。")

    add_heading(doc, "9.3 Target Credit机制", 2)
    add_mechanism(doc, "Target Credit申请、占用与释放",
                  "请求到达Target网络接口，Target服务模型明确可以承诺命令/数据/上下文资源。",
                  "按Target、读写类型和可选Bank占用Credit；必要时同时占用服务队列表项。",
                  "Credit递减并在事务上下文记录credit_held；服务完成路径只允许执行一次release。",
                  "Credit为0、服务队列满、Bank冲突门控、数据缓冲不足或Target处于暂停状态。",
                  "在已确认的‘Target真实占用结束点’返还；默认不等待Master最终接收，除非响应数据仍占用同一资源。",
                  "多响应读、写Grant后无数据、Target错误、超时、取消与重放可能导致重复或遗漏释放，必须以credit_held不变量保护。",
                  "Credit决定Target并发和队列位置；过早释放高估吞吐，过晚释放低估并行性，并可能错误地把响应拥塞表现为Target容量不足。")

    add_heading(doc, "9.4 FIFO、持续vld与自动重试", 2)
    add_mechanism(doc, "FIFO反压与事件唤醒机制",
                  "上游队头vld，调度器评估下游ready、Credit和Link可用性；或任一相关状态从不可用转为可用。",
                  "占用上游FIFO表项；握手成功时原子占用下游FIFO/Link/Target资源。阻塞期间不额外复制payload。",
                  "成功则下游入队且上游出队；失败则队头、vld和payload保持不变，并登记唯一主阻塞原因。",
                  "下游FIFO满、接口ready=0、Link序列化窗口未到、Credit为0或仲裁失败。",
                  "上游FIFO仅在下游成功接受后释放；阻塞本身不释放或重建事务。",
                  "下游由满变非满、多个状态同周期恢复、长时间vld、队头被取消、复位/排空和无状态事件的时间门限。",
                  "自动重试避免独立retry协议和重复包；必须确保所有可用性变化都发布唤醒，否则会形成‘包原地保留但永不前进’的模型死停。")

    add_heading(doc, "9.5 反压传播与进度保证", 2)
    add_body(doc, "反压按依赖路径逐级传播：Target服务资源不足使Target网络接口请求FIFO停止出队；FIFO满使目的Router LOCAL出口不可接收；拥塞沿Link和上游Router输入扩散，最终使Master网络接口注入FIFO保持队头，必要时通过其ready阻止上游继续发出。响应路径使用同样语义。模型不需要把一个全局ready组合穿越所有模块，而通过有限队列和事件传播表达延迟的反压波。")
    add_body(doc, "进度保证包含三项：第一，所有导致ready变化的事件都能唤醒相邻发送方；第二，仲裁提供有限饥饿保护，终结响应具有不会被请求永久占满的逻辑队列或保留份额；第三，全局无进展监视器能够区分合法长Target服务与等待环。持续反压不是错误，但当反压解除后必须在有界的调度延迟内观察到一次重评估。")


def add_section_10(doc, names):
    add_heading(doc, "10 延迟、带宽、排队及瓶颈模型", 1)
    add_heading(doc, "10.1 延迟分解", 2)
    add_body(doc, "端到端读延迟定义为READ_REQ被模型接受到最后READ_RSP被上游接受的cycle差；写延迟定义为WRITE_REQ被接受到WRITE_COMPLETE被上游接受的cycle差。模型同时记录零负载结构延迟和负载相关增量。建议使用以下分解：Ttotal=Tsrc_admission+Tinject_queue+Trequest_path+Ttarget_queue+Ttarget_service+Tresponse_path+Tupstream_accept，其中每条路径进一步分为Router固定流水、输出仲裁、FIFO等待、Link传播和Link序列化。")
    add_body(doc, "固定流水延迟在对应模块接收并承诺处理时只计一次；拓扑延迟是逐跳固定项的求和，不另加独立‘拓扑惩罚’；序列化由Link根据协议段计费字节和每cycle字节数计算占用窗口，不能同时在Router或Target端再次限速同一物理链路；排队延迟是实际入队到成功出队的时间差，不通过额外固定参数近似。Target服务延迟从Target接受请求或满足数据依赖开始，到Target完成并生成响应为止。")
    add_figure(doc, 5, names[4],
               "建立端到端延迟的唯一记账边界，并把源端、Ring最小割、Target和响应路径四类瓶颈与统计证据对应。",
               "源端准入、注入仲裁、逐跳Router、逐跳Link、Target入口、Target服务和响应回程构成串联路径；下方四个框表示容量上界。",
               "上方请求箭头表示时延累积；控制虚线把瓶颈框连接到对应资源；响应回程作为独立路径项计入。",
               "蓝色为事务路径，灰色为统计映射，橙色框为可能的容量瓶颈。",
               "硬件中Router与Link流水分界、协议头计费、全双工/半双工带宽，以及Target响应生成时点。")

    add_heading(doc, "10.2 链路带宽与序列化", 2)
    add_body(doc, "每条有向Ring Link拥有width_bytes_per_cycle、propagation_latency_cycles、max_inflight_segments和next_free_cycle。协议段计费字节packet_bytes=header_bytes+payload_bytes+optional_overhead_bytes；序列化周期Tser=max(1,ceil(packet_bytes/width_bytes_per_cycle))。当本段在cycle t提交时，链路下一次可启动时间更新为max(t,next_free_cycle)+Tser；段到达时间由提交起点加传播/流水延迟确定。该模型允许流水化链路在前一段尚未到达下一跳时启动后续段，但总启动速率受Tser限制。")
    add_body(doc, "若请求、写数据和响应共享一条物理等效链路，它们必须竞争同一个next_free_cycle和字节预算；若硬件使用独立子网，则每个子网可以拥有独立状态，但总面积/带宽比较必须使用一致口径。双向Ring的CW和CCW默认是两条独立有向资源；若物理上共享SerDes或端口，则需增加共享父级调度器。")
    add_mechanism(doc, "链路序列化与在途容量机制",
                  "Router输出winner准备提交，且目标Link方向和Traffic Class确定。",
                  "占用Link启动窗口Tser、一个in-flight表项及下游预留/到达事件。",
                  "next_free_cycle单调前移；段进入IN_FLIGHT并在到达时尝试写入下游FIFO。",
                  "当前cycle早于next_free_cycle、in-flight达到上限、下游预留策略失败或共享父级调度器未授权。",
                  "序列化窗口按时间到期；in-flight表项仅在下游成功接受后释放。",
                  "packet_bytes为0、超大段、传播时间小于Tser、到达时下游满、不同Traffic Class共享和双向物理耦合。",
                  "宽度决定饱和吞吐，传播决定RTT和所需OSD；把传播与Tser相加为非流水占用会低估带宽，把两者都当带宽限制会重复计算。")

    add_heading(doc, "10.3 Router与Target延迟", 2)
    add_body(doc, "Router固定延迟表示输入被接受到可参与输出提交的最少流水周期；输出仲裁等待和下游阻塞单独统计。Router可以配置cut-through等效模式，但若协议段仍以包级仲裁，则不得宣称模拟真实flit旁路。Target网络接口只计入口/出口整形和本地FIFO等待，不重复计Target内部服务时间。")
    add_body(doc, "Target服务模型至少具有base_latency_cycles、service_width_bytes_per_cycle、max_outstanding、request_queue_depth和读写独立/共享关系；可选增加Bank数、地址到Bank映射、row-hit/row-miss分布或经验延迟分布。L2和DDR不得共用一个未校准的固定延迟默认值。对正式性能运行，Target关键参数应显式配置且带来源；未提供时模型应拒绝运行或明确标识为non-signoff占位模式。")

    add_heading(doc, "10.4 热点和关键瓶颈形成", 2)
    add_body(doc, "热点可以由地址分布、interleave粒度、端点布点和路由方向共同形成。多个Master访问同一Target时，流量在靠近该Target的Ring Link和Target网络接口汇聚；即使Target总带宽足够，路径上的最小割也可能先饱和。反之，均匀分布到多个Target的流量仍可能因Ring布点不均或响应数据量较大而在某个方向形成回程热点。模型必须按链路、方向、Traffic Class和Target统计，而不能只报告全局平均利用率。")
    add_body(doc, "有效带宽上界为min(BWsource_injection, BWring_min_cut, BWtarget_aggregate, BWrsp_path, OSD×useful_payload/average_RTT)。其中BWtarget_aggregate仅在地址分布与Target并行度允许时可相加；BWring_min_cut必须针对实际流量矩阵求和；OSD项体现带宽时延积。瓶颈判定采用证据链而不是单一高利用率：资源利用率接近上限、上游队列或等待时间持续增长、放宽该资源能显著改善KPI，三者同时成立才认定为主瓶颈。")

    add_heading(doc, "10.5 防止延迟和带宽重复计算", 2)
    add_table(doc, ["物理/逻辑效应", "唯一记账模块", "禁止的重复表达", "审计统计"], [
        ("Router固定流水", "Ring Router", "Link传播或全局bus_latency再次叠加", "router_fixed_cycles"),
        ("输出仲裁等待", "Ring Router", "以随机拥塞惩罚再加一次", "arb_wait_hist/max"),
        ("Link传播", "Ring Link", "按hop另加拓扑惩罚", "link_prop_cycles"),
        ("Link序列化", "Ring Link", "Router端口和Target端口同时按同宽度限速", "link_busy_cycles/bytes"),
        ("FIFO排队", "拥有该FIFO的模块", "固定queue_penalty", "queue_wait/occupancy"),
        ("Target服务", "Target服务模型", "Target网络接口固定响应延迟重复叠加", "service_latency/busy"),
        ("OSD/Credit阻塞", "对应资源所有者", "同时作为额外固定延迟", "stall_cycles_by_reason"),
    ], widths=[1900,1850,3500,2110], font_size=8.2, first_col_bold=True)


def add_section_11(doc):
    add_heading(doc, "11 配置和统计设计", 1)
    add_heading(doc, "11.1 配置原则与生效链", 2)
    add_body(doc, "配置采用版本化Schema，按全局、拓扑、Master、Router/Link、Target和运行控制分层。加载顺序建议为‘Schema默认→平台基线→端点类型模板→实例覆盖→场景覆盖’，每个最终值保留来源、单位和覆盖路径。模型启动前生成规范化配置快照及指纹，并输出‘参数→资源所有者→状态变量/调度器→预期统计’的生效清单；任何未知字段、单位冲突、地址重叠或逻辑冲突默认报错，不允许静默忽略。")
    add_body(doc, "建议默认值仅用于可运行的架构探索，不代表硬件事实。对Target服务延迟、带宽、物理通道共享和协议头开销等签核敏感参数，正式模式要求显式配置或从经批准的平台模板继承；缺失时模型应拒绝进入signoff运行。运行中原则上不允许修改影响事务语义或资源容量的参数，必要的动态频率、带宽节流或Target暂停应作为带时间戳的控制事件记录。")

    add_heading(doc, "11.2 参数配置表", 2)
    add_body(doc, "下表范围是模型建议支持范围，用于输入校验和容量规划；不是硬件承诺。‘建议默认’标为MANDATORY的项目在正式性能运行中必须显式给出。")

    add_heading(doc, "11.2.1 全局、拓扑与路由参数", 3)
    add_table(doc, ["参数", "定义/单位", "建议默认", "建议范围", "作用位置与性能影响"], [
        ("schema_version", "配置Schema版本/字符串", "1.0", "受版本管理", "配置与约束管理；决定字段解释和兼容策略。"),
        ("clock_frequency", "模型时钟/GHz", "1.0（探索）", "0.2–4.0", "仅用于cycle与时间/带宽换算；不改变以cycle配置的事件顺序。"),
        ("topology_type", "互连拓扑/枚举", "BIDIR_RING", "BUS/RING/XBAR/MESH", "选择传输平面；端点事务接口不变。"),
        ("master_count", "Master实例数/个", "4（探索）", "1–64", "决定注入端数量、ID宽度和竞争规模。"),
        ("target_count", "Target实例数/个", "4（探索）", "1–32", "决定地址分布、服务并行度和节点规模。"),
        ("ring_node_count", "Ring节点数/个", "自动≥端点布点", "2–128", "决定跳数、链路数量和平均RTT。"),
        ("endpoint_placement", "端点到节点映射/表", "均匀分散（探索）", "完整显式映射", "直接影响路径重叠和热点；正式运行建议显式。"),
        ("route_policy", "路径策略/枚举", "SHORTEST_HASH_TIE", "CW/CCW/SHORTEST_*", "影响方向利用率、顺序与可复现性。"),
        ("interleave_mode", "交织策略/枚举", "LINEAR", "NONE/LINEAR/XOR", "决定Target负载和地址局部性。"),
        ("interleave_granule", "交织粒度/Byte", "128（探索）", "64–4096，2的幂", "小粒度提高均衡但增加分段；必须≥最小传输对齐。"),
        ("address_map", "地址域/结构", "MANDATORY", "不重叠、完整定义", "地址映射与交织引擎；非法/空洞地址默认报错。"),
    ], widths=[1650,1900,1450,1650,2710], font_size=7.7, first_col_bold=True)

    add_heading(doc, "11.2.2 Master、Router与Link参数", 3)
    add_table(doc, ["参数", "定义/单位", "建议默认", "建议范围", "作用位置与性能影响"], [
        ("master_read_osd", "每Master读Outstanding/事务", "16（探索）", "1–256", "Master网络接口；过小限制带宽时延积，过大增加拥塞。"),
        ("master_write_osd", "每Master写Outstanding/事务", "16（探索）", "1–256", "覆盖请求至最终完成的写生命周期。"),
        ("master_req_fifo_depth", "请求FIFO/段", "8", "1–256", "吸收源端突发；不替代OSD。"),
        ("master_wdata_fifo_depth", "写数据FIFO/段", "8", "1–256", "影响Grant后数据供应和头数据解耦。"),
        ("router_input_fifo_depth", "每端口每逻辑类/段", "4", "1–64", "决定局部突发吸收和HOL；过深会增加缓冲延迟。"),
        ("router_pipeline_latency", "Router固定流水/cycle", "1（探索）", "0–8", "每跳唯一固定项；0仅用于理想敏感度。"),
        ("arbitration_policy", "输出仲裁/枚举", "RR_WITH_AGE", "RR/WRR/FIXED/AGE", "决定服务份额、公平性和尾延迟。"),
        ("traffic_class_weights", "流量类权重/无量纲", "1:1:1", "1–255/类", "共享链路时分配请求、写数据、响应份额。"),
        ("starvation_threshold", "年龄逃逸阈值/cycle", "1024（探索）", "16–1,000,000", "限制最长等待；过低会破坏QoS权重。"),
        ("link_width", "每有向链路宽度/Byte per cycle", "32（探索）", "8–256", "Ring Link序列化吞吐的唯一宽度参数。"),
        ("link_propagation_latency", "每跳传播/cycle", "1（探索）", "1–16", "影响零负载RTT和所需OSD，不直接降低流水吞吐。"),
        ("link_max_inflight", "每链路在途段/段", "16", "1–512", "限制传播管线容量；应覆盖带宽×传播时间。"),
        ("channel_share_mode", "三类流量资源关系/枚举", "SHARED", "SHARED/SPLIT_REQ_RSP/FULL_SPLIT", "P0参数；直接决定读写互扰和总吞吐。"),
    ], widths=[1650,1900,1450,1650,2710], font_size=7.45, first_col_bold=True)

    add_heading(doc, "11.2.3 Target与协议参数", 3)
    add_table(doc, ["参数", "定义/单位", "建议默认", "建议范围", "作用位置与性能影响"], [
        ("target_credit", "每Target可占用上下文/事务", "MANDATORY", "1–1024", "Target网络接口；限制并发准入。"),
        ("target_req_fifo_depth", "Target入口FIFO/段", "8（探索）", "1–512", "吸收Ring突发并形成可见排队。"),
        ("target_rsp_fifo_depth", "Target响应FIFO/段", "8（探索）", "1–512", "吸收回程阻塞；决定Credit是否需保留。"),
        ("target_base_latency", "基础服务延迟/cycle", "MANDATORY", "≥0", "Target服务模型；L2/DDR必须分别配置。"),
        ("target_service_width", "有效服务宽度/Byte per cycle", "MANDATORY", ">0", "Target长期吞吐上限；不得与Ring宽度混用。"),
        ("target_max_outstanding", "Target内部并发/事务", "MANDATORY", "1–4096", "服务模型内部并发，可能与外部Credit相等或更小。"),
        ("target_bank_count", "可选Bank数/个", "1（抽象）", "1–256", "用于表达Bank并行和冲突；未校准时保持1。"),
        ("target_latency_distribution", "服务延迟分布/结构", "FIXED（探索）", "FIXED/TABLE/TRACE", "用于尾延迟和动态服务能力；需参考数据。"),
        ("request_header_bytes", "请求头计费/Byte", "MANDATORY", "1–256", "影响READ_REQ/WRITE_REQ链路序列化。"),
        ("response_header_bytes", "响应头计费/Byte", "MANDATORY", "1–256", "影响Grant、完成和错误响应。"),
        ("data_beat_bytes", "数据段粒度/Byte", "64（探索）", "8–1024", "影响分段、仲裁粒度和响应交织。"),
        ("write_protocol_mode", "写协议/枚举", "GRANT_THEN_DATA", "GRANT/POSTED/PARALLEL", "P0参数；决定数据缓冲、RTT和完成语义。"),
        ("credit_release_point", "Credit释放/枚举", "TARGET_RESOURCE_FREE", "SERVICE_DONE/RSP_ENQ/RSP_ACCEPT", "P0参数；必须匹配真实资源生命周期。"),
    ], widths=[1650,1900,1450,1650,2710], font_size=7.35, first_col_bold=True)

    add_heading(doc, "11.2.4 运行、检查与统计参数", 3)
    add_table(doc, ["参数", "定义/单位", "建议默认", "建议范围", "作用位置与性能影响"], [
        ("warmup_cycles", "预热时间/cycle", "按场景显式", "0–仿真长度", "预热不计入主KPI，但保留稳定性检查。"),
        ("measurement_cycles", "测量窗口/cycle", "MANDATORY", ">0", "所有吞吐/利用率使用同一窗口。"),
        ("random_seed", "随机种子/整数", "1", "32/64位", "分布流量和Target随机延迟可复现。"),
        ("deadlock_watchdog_cycles", "无进展阈值/cycle", "100000（探索）", "大于最大合法服务延迟", "检测等待环；不得误判合法长延迟。"),
        ("trace_level", "追踪级别/枚举", "SUMMARY", "OFF/SUMMARY/TRANSACTION/VERBOSE", "影响可诊断性和仿真速度。"),
        ("latency_histogram_bins", "直方图边界/cycle数组", "对数分桶", "覆盖p99以上", "计算分位数并避免仅报告平均。"),
        ("stats_interval", "周期统计窗口/cycle", "10000（探索）", "100–测量窗口", "观察拥塞随时间变化；过小增加开销。"),
        ("strict_config", "非法配置处理/布尔", "true", "true/false", "正式模式必须true；禁止静默忽略参数。"),
        ("signoff_mode", "签核运行/布尔", "false", "true/false", "true时要求所有MANDATORY参数有批准来源。"),
    ], widths=[1650,1900,1450,1650,2710], font_size=7.7, first_col_bold=True)

    add_heading(doc, "11.3 统计设计与瓶颈归因", 2)
    add_body(doc, "统计按全局、Master、Target、Router输出、Link方向和Traffic Class分层。吞吐同时报告完成事务数、有效payload字节、链路计费字节和byte/cycle；延迟至少报告count、mean、min、max、p50、p90、p95和p99，并保留分段等待分解。队列统计包括占用时间积分、占用率直方图、满周期、入出队数和队头最大等待。仲裁统计包括候选数、grant、loser等待、最大等待和每Master服务份额。")
    add_body(doc, "每次发送失败只登记一个主阻塞原因，按优先级选择‘OSD/Credit不足→下游FIFO满→Link忙/在途满→仲裁失败→顺序依赖→Target服务门控→上游不接收’，同时允许记录辅助位图。主原因计数用于占比相加，辅助位图用于分析共现；否则同一cycle多个原因重复计数会产生超过100%的停顿解释。瓶颈报告需要把利用率、等待增长和参数放宽实验关联起来。")
    add_table(doc, ["对象", "必备统计", "主要判断"], [
        ("Master", "accepted/completed、读写OSD占用、注入stall、延迟分位数、有效带宽", "源端供给、OSD/BDP不足、Master间公平性"),
        ("Router输出", "候选数、grant、按输入等待、最大年龄、Traffic Class份额", "竞争位置、饥饿、QoS执行"),
        ("Ring Link", "计费字节、busy cycle、in-flight、下游阻塞、按类利用率", "最小割、方向热点、响应带宽不足"),
        ("Target网络接口", "Credit轨迹、FIFO占用、准入stall、响应回注stall", "入口容量、Credit释放和回程阻塞"),
        ("Target服务模型", "queue/service latency、busy、并发、Bank冲突、完成率", "Target真实服务瓶颈"),
        ("全局", "总吞吐、延迟、Jain指数、无进展周期、事务/资源不变量", "系统拐点、正确性和可解释性"),
    ], widths=[1700,4600,3060], font_size=8.2, first_col_bold=True)

    add_heading(doc, "11.4 参数有效性与可追溯性", 2)
    add_body(doc, "每个配置参数必须在启动日志中打印最终值、来源和消费者，并在统计文件中携带配置指纹、模型版本、随机种子、工作负载版本和测量窗口。自动化参数审计至少验证：减小FIFO不会提高无拥塞理论带宽；增大Link宽度不应改变纯传播周期；增大传播延迟在OSD充足时不改变稳态链路峰值；增大Target Credit只在Target并发受限时改善吞吐；改变interleave能改变Target分布但不改变总请求字节。若方向不符合物理预期，视为参数未生效或重复建模。")


def add_section_12(doc, names):
    add_heading(doc, "12 精度定义、校准和验证方案", 1)
    add_heading(doc, "12.1 80%精度的可验证定义", 2)
    add_body(doc, "80%不是‘大多数曲线看起来相似’，也不是单一平均吞吐的拟合值。对每个场景s和KPI k，定义归一化准确度A(s,k)=max(0,1-|Model-Reference|/max(|Reference|,epsilon_k))。epsilon_k用于接近0的统计，必须在验收计划中预先给出。场景得分为KPI加权和，整体得分为场景加权和。建议权重为：稳态吞吐35%，p50/p95延迟25%，饱和拐点15%，多核/Target扩展趋势15%，瓶颈分类与stall归因10%。")
    add_body(doc, "目标通过条件是留出验收集的整体加权准确度不低于0.80，并同时满足硬门限。建议硬门限：功能与资源不变量100%通过；至少80%的验收场景吞吐相对误差不超过15%；至少80%的场景p50和p95延迟相对误差不超过20%；所有规模/OSD/带宽扫描的单调方向正确；主瓶颈分类正确率不低于80%；任何P0代表场景的关键KPI误差不得超过30%，除非评审批准为模型适用域外。权重和门限是建议值，需专家签核后冻结。")
    add_callout(doc, "精度声明边界", "在参考平台、配置映射、场景集和验收报告未完成前，文档和模型只能写‘精度目标约80%’，不得写‘精度达到80%’。通过验收后也必须同时声明适用端点规模、流量类型、Target模型和参数范围。")

    add_heading(doc, "12.2 基准、数据集与测量对齐", 2)
    add_body(doc, "参考数据优先级为经验证的硅上性能计数器与trace、已收敛的RTL性能仿真、经批准的高保真平台模型。每个参考点必须具有同一地址映射、端点布点、时钟、协议头开销、读写比例、请求长度、OSD和Target配置。预热、测量窗口、随机种子和统计起止点必须一致；如果参考平台无法观测某个内部等待项，则该项只用于模型内部归因，不作为独立签核KPI。")
    add_body(doc, "场景集按目的分为单元表征、校准集和留出验收集。单元表征用于测量零负载跳延迟、单链路峰值、单Target服务曲线和OSD带宽时延积；校准集覆盖代表性均匀、热点、读写混合和背压条件，用于调整有硬件含义的参数；留出集在参数冻结前不用于调参，必须包含未见过的流量矩阵、端点规模或长度组合，以检测过拟合。")
    add_figure(doc, 6, names[5],
               "规定从参考数据到模型签核的闭环，防止在同一数据上反复调参后宣称精度达标。",
               "基准定义、场景与数据集、模型执行、KPI对齐、误差分析、参数校准、版本冻结、留出集验收和评审签核依次构成闭环。",
               "灰色控制箭头表示数据和配置流；红色回退表示结构或参数误差的回查；不存在留出集直接反馈参数的路径。",
               "灰色为流程控制，红色为不通过回退，蓝色为模型冻结，绿色为验收。",
               "参考平台优先级、KPI权重、epsilon、硬门限、留出集规模及异常豁免权限。")

    add_heading(doc, "12.3 校准顺序", 2)
    add_body(doc, "校准按可辨识性从低负载到高负载进行。第一阶段用单Master、单Target和最短/最长路径锁定接口、Router、Link和Target的固定延迟；第二阶段用持续大包流锁定链路和Target峰值带宽及序列化开销；第三阶段扫描OSD和Credit锁定带宽时延积与并发上限；第四阶段引入多Master、热点和读写混合校准仲裁权重、队列深度和Target动态服务；最后用未参与调参的场景验收。不得在固定延迟未锁定前用队列惩罚拟合低负载，也不得用Target延迟补偿错误的链路共享关系。")
    add_table(doc, ["阶段", "主要输入", "校准参数", "通过观察量", "典型误差归因"], [
        ("结构一致性", "端点图、地址图、协议定义", "拓扑、映射、路径、通道共享", "路径/字节/事务计数一致", "结构假设错误"),
        ("零负载", "单事务、不同跳数", "接口/Router/Link/Target固定延迟", "逐段延迟线性可解释", "延迟边界或起止点错误"),
        ("峰值带宽", "持续流、OSD充足", "Link/Target宽度、头开销", "平台期望峰值与利用率", "重复限速或计费字节错误"),
        ("并发", "OSD/Credit扫描", "并发深度、in-flight", "拐点和BDP", "申请/释放点错误"),
        ("竞争", "多Master热点/混合读写", "仲裁、FIFO、Target动态服务", "份额、尾延迟、热点", "公平性、HOL或服务抽象不足"),
        ("留出验收", "未见场景", "无；参数已冻结", "加权得分与硬门限", "过拟合或适用域不足"),
    ], widths=[1200,2100,2100,2100,1860], font_size=8.0, first_col_bold=True)

    add_heading(doc, "12.4 验证矩阵", 2)
    add_body(doc, "验证按机制覆盖组织，不以某个临时用例名称定义架构能力。基础正确性覆盖单Master读写、多响应、写授权、跨边界分段和错误终结；竞争覆盖多个Master争同一输出、同一Link和同一Target；映射覆盖均匀、热点、规则步长和XOR散列；流控覆盖每一级FIFO满、OSD耗尽、Credit耗尽、Link在途满、Target长期不接收及恢复；性能覆盖零负载、逐级加压、饱和平台、OSD/带宽/延迟/深度扫描；鲁棒性覆盖复位排空、ID回绕、最大长度、连续反压和无进展检测。")
    add_table(doc, ["验证域", "关键激励", "主要通过条件"], [
        ("事务闭环", "读/写、多段、多响应、错误", "TxnKey关联正确，Expected/Received闭合，资源回零"),
        ("真实竞争", "多Master同周期访问共享输出/Target", "观察到显式仲裁、等待和可重复服务份额"),
        ("反压恢复", "逐级填满后解除下游阻塞", "包原地保持、无丢包，状态变化后自动重试并最终完成"),
        ("死锁/饥饿", "响应拥塞、混合类、长期热点", "无等待环；最大等待满足策略边界或被watchdog捕获"),
        ("参数生效", "单参数扫全范围与边界值", "KPI方向、拐点和统计作用点符合预期"),
        ("精度", "校准集与留出集", "整体≥0.80且全部硬门限通过"),
        ("仿真性能", "大端点、长时间、高拥塞", "事件数、内存和运行时间满足项目预算，统计无抽样偏差"),
    ], widths=[1600,3900,3860], font_size=8.4, first_col_bold=True)

    add_heading(doc, "12.5 签核输出", 2)
    add_body(doc, "每次精度签核输出应包含：模型和配置指纹、参考平台版本、完整场景表、原始KPI、相对/绝对误差、加权得分、硬门限结果、瓶颈混淆矩阵、异常场景归因、参数变更记录和适用域声明。任何手工豁免必须记录批准人、原因、预计影响和后续关闭计划。模型升级后若改变事务语义、通道共享、计费字节或Target服务结构，必须重新执行全量留出集验收。")


def add_section_13(doc):
    add_heading(doc, "13 替代方案与设计权衡", 1)
    add_heading(doc, "13.1 拓扑方案比较", 2)
    add_table(doc, ["方案", "主要收益", "主要代价", "适用性判断"], [
        ("共享总线", "模型简单、单一仲裁点、低负载延迟易校准。", "所有端点共享带宽，无法表达空间并行；端点增长后扩展性差。", "适合小规模、低带宽、确有单一共享媒介的系统或早期上界模型。"),
        ("Crossbar", "低跳数、不同Target间并行，路径直观。", "输入×输出仲裁和面积随端点扩展；物理布局/长线代价不易抽象。", "适合端点数有限、性能要求高且硬件接近全连接交换的系统。"),
        ("双向Ring", "结构规则、链路局部、可表达多跳与空间复用；建模复杂度中等。", "平均跳数随规模增长，热点链路和阻塞传播明显；最小割有限。", "适合中等规模端点、布局呈环或可用Ring近似、需要研究路径与热点的首版。"),
        ("Mesh", "更高扩展性和路径多样性，适合大规模分布端点。", "路由/VC/死锁和校准复杂度高；参数和统计数量显著增加。", "当端点规模、二维布局或多流量域使Ring最小割持续不足时作为演进方案。"),
        ("分层/混合", "可按簇聚合，平衡局部带宽与全局成本。", "层间接口和瓶颈归因更复杂，配置空间增大。", "适合Core簇内Crossbar、簇间Ring/Mesh等已有硬件分层。"),
    ], widths=[1300,2650,2800,2610], font_size=8.3, first_col_bold=True)

    add_heading(doc, "13.2 Ring基线的收益与代价", 2)
    add_body(doc, "选择Ring的主要收益是能以较少状态表达硬件架构师关心的逐跳距离、方向争用、热点最小割和下游反压，同时允许不同链路并行推进。相较共享总线，它能真实反映端点布点和路径重叠；相较Crossbar，它避免假定任意输入到任意输出都具有独立高速通路；相较Mesh，它不需要在首版引入二维路由和VC死锁规避。")
    add_body(doc, "其代价是大规模端点下平均/最坏跳数上升，少数热点Target会使相邻链路先饱和，包级仲裁不能复现flit交错和VC级HOL，双向最短路也可能与真实物理拓扑不完全一致。因此，Ring只应在端点数、流量矩阵和硬件布局近似合理的适用域内签核。建议以‘最坏链路利用率、平均/最大跳数、热点Target相邻链路带宽裕量和RTL/硅对标误差’作为Ring继续适用的门槛，而不是用固定Core数量一刀切。")

    add_heading(doc, "13.3 建模粒度权衡", 2)
    add_table(doc, ["粒度", "能表达的行为", "省略的行为", "建议"], [
        ("事务固定延迟", "功能路径和理想RTT", "竞争、序列化、反压、热点", "仅作为sanity上界，不作为本项目基线"),
        ("协议段/包级（本方案）", "多跳、带宽、分流、输出竞争、有限资源", "flit交错、VC分配、逐拍crossbar", "首版基线；达到约80%目标的成本/收益平衡点"),
        ("flit/VC级", "细粒度HOL、credit返回、router微流水", "物理实现以下细节", "仅在对标误差证明包级不足时升级"),
        ("RTL逐拍", "协议与微架构周期行为", "系统规模与运行速度受限", "用于参考和局部签核，不替代快速ESL探索"),
    ], widths=[1800,3000,2700,1860], font_size=8.5, first_col_bold=True)

    add_heading(doc, "13.4 演进路径", 2)
    add_body(doc, "方案通过稳定的Master/Target事务接口和独立传输平面支持演进。若需要更高精度，可按证据逐步增加：共享链路内的多VC、flit分段与wormhole占用、自适应路由、分层时钟、QoS带宽保证、DDR命令级Target或一致性Traffic Class。每次升级必须说明解决的已观测误差、增加的状态与运行成本、对旧配置的兼容性，以及新的死锁和验证义务。")


def add_section_14(doc):
    add_heading(doc, "14 风险、假设和待确认事项", 1)
    add_heading(doc, "14.1 关键风险", 2)
    add_table(doc, ["风险", "概率/影响", "触发信号", "缓解与关闭标准"], [
        ("通道共享假设错误", "高/高", "混合读写峰值与参考差异大，双向总带宽异常", "P0冻结物理等效共享关系；单类与混合流量共同对标。"),
        ("OSD/Credit释放错误", "中/极高", "计数不回零、吞吐异常随响应阻塞变化", "资源所有权不变量、故障注入、长反压后排空。"),
        ("写请求与数据关联不完整", "中/极高", "DataTag错配、Grant后永久等待、跨事务数据污染", "冻结写协议状态机；beat/超时/错误覆盖。"),
        ("响应依赖导致死锁", "中/极高", "存在在途事务但全局长期无握手", "逻辑类分队列、响应保留进度、等待图与watchdog。"),
        ("仲裁受软件事件顺序影响", "中/高", "改变实例创建顺序即改变Master带宽", "候选收集+原子winner提交；确定性回归。"),
        ("延迟或带宽重复计算", "高/高", "单参数扫描出现非物理方向或峰值过低", "唯一资源所有者矩阵和分析上界对照。"),
        ("Target模型过度简化", "高/高", "低负载准确但热点/尾延迟失配", "分阶段校准；需要时替换为经验分布或高保真模型。"),
        ("参数存在但未生效", "中/高", "改变配置KPI和内部统计均不变", "配置消费者清单、严格未知字段、敏感度门禁。"),
        ("统计不足或重复归因", "中/高", "stall占比>100%或无法定位饱和点", "主原因+辅助位图；分层计数和因果放宽实验。"),
        ("对标数据不足", "高/极高", "80%只能在校准场景成立或无留出集", "评审前确定数据Owner、参考版本和采集计划。"),
        ("仿真运行成本失控", "中/中", "高拥塞场景事件数/内存超预算", "事件去重、按需trace、统计采样和规模基准。"),
    ], widths=[2200,1200,2800,3160], font_size=7.7, first_col_bold=True)

    add_heading(doc, "14.2 假设集中清单", 2)
    add_body(doc, "本稿采用的保守假设包括：首版以单播读写为主；统一cycle时间基准；协议段为最小仲裁单位；双向Ring使用确定性最短路；请求、写数据和响应独立逻辑排队；三类流量默认共享每方向链路带宽；写采用Grant后发数据；Target Credit在真实Target资源结束时释放；地址跨交织边界时先分段；非法地址不回退到默认Target；错误通过终结响应闭环；正式运行的Target服务和协议开销参数必须显式给出。")
    add_body(doc, "这些假设共同影响吞吐上界、排序、死锁依赖和精度适用域。任何一项被评审推翻，都必须更新事务状态机、资源生命周期、参数Schema、验证矩阵和相关架构图，而不能只修改一个默认参数。")

    add_heading(doc, "14.3 独立反对者审查结果", 2)
    add_body(doc, "初稿完成后按资深硬件架构师的反对者视角执行独立审查。下表列出问题、严重度、对方案的挑战以及修订状态。‘已纳入’表示正文已给出防护或验证要求，但仍可能需要专家提供硬件事实；‘待决议’表示没有足够输入，不能在文档中虚构结论。")

    add_heading(doc, "Critical", 3)
    add_table(doc, ["编号", "问题", "后果", "修改建议/状态"], [
        ("C-01", "请求、写数据、响应的物理等效共享关系未有硬件依据。", "可能把可用带宽扩大2–3倍，且读写互扰错误。", "正文将逻辑分流与物理共享解耦，默认共享；列为P0待决议。"),
        ("C-02", "多响应读与两阶段写的OSD/Credit释放点若不冻结，事务可能提前完成、泄漏或过释放。", "功能错误并直接扭曲吞吐。", "增加Expected/Received、不变量和分层释放规则；Credit真实释放点仍为P0。"),
        ("C-03", "顺序域、原子/屏障和错误终结范围未定义。", "模型结果可能不满足上层可见语义，性能也可能因错误重排被高估。", "基线明确仅局部FIFO顺序，其他均不宣称支持；列为P0。"),
        ("C-04", "80%目标尚无参考平台、场景权重和留出集。", "目标不可证实，任何准确度声明均不成立。", "建立量化评分与校准流程；参考Owner和数据计划仍为P0。"),
    ], widths=[900,3600,2200,2660], font_size=8.0, first_col_bold=True)

    add_heading(doc, "Major", 3)
    add_table(doc, ["编号", "问题", "后果", "修改建议/状态"], [
        ("M-01", "输出端winner若受事件调用顺序决定，不是明确仲裁。", "多核竞争存在但公平性不可复现。", "规定候选收集、原子提交和RR+年龄机制；详细设计必须落地。"),
        ("M-02", "响应流量可能被请求/写数据占满共享资源。", "形成协议依赖死锁或极端尾延迟。", "独立逻辑队列、响应进度份额和等待图检查已纳入。"),
        ("M-03", "写Grant后无数据、部分数据或错误缺少闭环。", "Target上下文和Credit永久占用。", "增加DataBeat、超时/错误终结和排空要求；超时数值待协议确认。"),
        ("M-04", "固定延迟、hop惩罚和序列化可能重复记账。", "零负载和峰值同时失真，无法校准。", "增加唯一记账矩阵和分阶段校准顺序。"),
        ("M-05", "仅有全局统计无法证明热点和主瓶颈。", "性能结果不可解释，参数调整可能误导。", "增加每Master/输出/Link/Target/Traffic Class统计和主阻塞原因。"),
        ("M-06", "配置参数可能被解析但未连接到状态机。", "看似可配置，实际无性能影响。", "增加消费者清单、配置指纹和单参数敏感度门禁。"),
        ("M-07", "Target只用固定延迟可能无法覆盖DDR尾延迟和Bank冲突。", "热点与混合流量误差超出预算。", "Target接口支持分布/Bank扩展，以对标证据决定是否启用。"),
        ("M-08", "下游恢复若未产生唤醒事件，持续vld包可能永久滞留。", "非协议死锁，仿真无法排空。", "明确所有ready/Credit/时间窗变化触发重评估，并加入恢复有界性验证。"),
    ], widths=[900,3600,2200,2660], font_size=7.8, first_col_bold=True)

    add_heading(doc, "Minor", 3)
    add_table(doc, ["编号", "问题", "修改建议/状态"], [
        ("m-01", "未对齐和跨交织边界请求可能被隐式送往单一Target。", "规定源端分段和父事务聚合，并加入边界验证。"),
        ("m-02", "等距固定方向会形成可重复但非真实的方向偏置。", "默认使用TxnKey稳定散列，保留固定方向对照。"),
        ("m-03", "阻塞次数与阻塞周期可能混用。", "分别统计事件数、持续周期和队头等待，报告中标注口径。"),
        ("m-04", "预热和测量窗口不一致会污染对标。", "运行配置中强制记录窗口，并将所有吞吐统计绑定同一窗口。"),
        ("m-05", "ID回绕和Epoch省略条件不明确。", "仅当LocalTxnID空间保证覆盖最大在途生命周期时允许无Epoch。"),
        ("m-06", "复位和排空语义未固化。", "要求停止新注入、等待在途闭环或生成终结错误，并验证idle全空。"),
    ], widths=[1000,4100,4260], font_size=8.4, first_col_bold=True)

    add_heading(doc, "Suggestion", 3)
    add_table(doc, ["编号", "建议", "预期价值"], [
        ("S-01", "提供自动敏感度与理论上界对照报告。", "及早发现死参数、重复限速和不可辨识参数。"),
        ("S-02", "配置Schema、统计Schema和图中术语使用同一机器可读词典。", "减少文档、配置和实现命名漂移。"),
        ("S-03", "为每个事务保留可抽样的资源时间线。", "可重建p99延迟和多级反压因果链。"),
        ("S-04", "建立Shared Bus/Ring/Crossbar的同口径理想基线。", "区分拓扑收益与Target/流量变化。"),
        ("S-05", "对配置和模型版本生成不可变指纹。", "保证对标结果可复现、可审计。"),
    ], widths=[1000,5000,3360], font_size=8.5, first_col_bold=True)

    add_heading(doc, "14.4 审查闭环结论", 2)
    add_body(doc, "独立审查未发现无法通过架构设计解决的事务断链，但发现四项需要外部硬件/项目决策的Critical事项。正文已经补充了TxnKey、ExpectedRsp、写DataTag、资源权威所有者、持续vld自动重试、唯一延迟/带宽记账、分层统计和量化精度流程。剩余Critical不是文字缺失，而是必须由评审专家确认的产品和硬件事实，因此不得在方案中自行假定关闭。")


def add_section_15(doc):
    add_heading(doc, "15 专家评审议题与阶段结论", 1)
    add_heading(doc, "15.1 需要拍板的关键问题", 2)
    add_table(doc, ["决策ID", "专家需拍板的问题", "建议结论", "决议输出"], [
        ("D-01", "首版是否以事务级双向Ring作为正式模型基线", "通过；保持端点/Target接口拓扑无关", "拓扑、端点规模适用域和演进条件"),
        ("D-02", "三类Traffic Class的物理等效资源关系", "默认共享每方向Link带宽，逻辑分队列", "共享矩阵、宽度、仲裁和保留份额"),
        ("D-03", "写请求、Grant、写数据和完成协议", "首版采用Grant后发数据", "DataTag、beat、缓冲、错误/超时状态机"),
        ("D-04", "读写顺序域、同地址hazard、atomic/barrier范围", "首版只保证局部FIFO与TxnKey归属", "支持范围、顺序点和非目标清单"),
        ("D-05", "Master OSD与Target Credit的申请/释放点", "OSD到上游终结；Credit到Target真实资源结束", "每种事务的资源生命周期表"),
        ("D-06", "仲裁、QoS和饥饿边界", "RR+年龄逃逸；按需扩展WRR", "权重、响应进度份额、最大等待"),
        ("D-07", "Target服务模型保真层次", "延迟+带宽+并发为基线，Bank/分布按证据启用", "L2/DDR模板及参数来源Owner"),
        ("D-08", "80%精度验收定义", "采用加权KPI+硬门限+留出集", "参考平台、场景/权重、epsilon和豁免权"),
        ("D-09", "详细设计与编码准入", "Critical关闭后有条件通过", "评审纪要、Owner、完成日期和基线版本"),
    ], widths=[900,3350,3350,1760], font_size=7.8, first_col_bold=True)

    add_heading(doc, "15.2 进入详细设计和编码的门禁", 2)
    add_body(doc, "建议将准入分为架构闭环、详细设计和全面编码三层。架构闭环要求D-01至D-08形成书面决议，尤其是通道共享、顺序、写协议、Credit释放和精度基准；详细设计准入要求模块接口、事务状态机、资源生命周期、事件阶段顺序、配置Schema和统计Schema完成评审；全面编码准入还要求Critical/Major验证项具有可执行测试计划、Target基线参数有Owner、对标数据采集计划已承诺。")
    add_table(doc, ["门禁", "必须满足", "未满足时允许的工作"], [
        ("G0 架构决议", "D-01至D-08有明确结论，Critical有Owner和关闭证据。", "文档完善、基准数据准备；不得冻结性能语义。"),
        ("G1 详细设计", "接口、状态机、资源表、调度顺序、Schema和验证矩阵评审通过。", "可开发配置/统计基础设施和非争议原型。"),
        ("G2 全面编码", "Major验证项可执行，Target参数/参考数据计划已落实。", "若未满足，不应以模型数值支持架构签核。"),
        ("G3 精度签核", "留出集加权准确度≥0.80且硬门限全通过。", "可继续校准或缩小适用域；不得宣称已达80%。"),
    ], widths=[1700,5000,2660], font_size=8.5, first_col_bold=True)

    add_heading(doc, "15.3 明确结论", 2)
    add_callout(doc, "评审建议：有条件具备进入详细设计，不具备无条件全面编码", "本文已达到方案级架构评审所需的边界、模块职责、事务闭环、资源生命周期、时延/带宽模型、配置/统计、校准与验证深度。若专家在本次评审中关闭D-02至D-05及D-08等Critical关联决策，并将决议写回基线，则方案可进入详细设计；在这些事项未关闭前，只建议开展不依赖其语义的基础设施和验证框架工作。模型在完成留出集验收前只能声明‘精度目标约80%’，不能声明已经达到。", fill="FFF7ED", accent=AMBER)
    add_body(doc, "评审通过后的首批详细设计输出应包括：接口信号与payload字段表、读写状态机、父子事务和多响应计数规则、每类资源的acquire/release时序、同cycle事件优先级、各输出仲裁伪代码、配置Schema、统计字段Schema、错误/复位/排空语义，以及覆盖本文件Critical/Major问题的验证计划。")

    add_heading(doc, "附录A 交付物索引", 1)
    add_table(doc, ["交付物", "位置/形式", "说明"], [
        ("完整设计文档", "本Word文档", "包含正式正文、一页摘要、参数表、风险、决策和阶段结论。"),
        ("一页专家评审摘要", "封面后第1页", "用于评审会预读和结论快速对齐。"),
        ("六张架构图", "正文图1–图6", "每图前含图纸设计说明。"),
        ("可编辑图源", "editable_diagrams/*.svg", "SVG 1.1，白色背景、正交连线、统一线型，可继续编辑。"),
        ("参数配置表", "第11.2节", "包含定义、单位、建议默认、范围、作用位置和性能影响。"),
        ("风险与独立审查", "第14节", "Critical/Major/Minor/Suggestion及修改建议。"),
        ("专家拍板与准入结论", "第15节", "D-01至D-09、G0至G3和明确阶段结论。"),
    ], widths=[1900,2700,4760], font_size=8.6, first_col_bold=True)


def build_document(names):
    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)
    props = doc.core_properties
    props.title = "多核SoC总线模型方案与架构设计文档"
    props.subject = "多核SoC片上互连ESL性能模型架构评审"
    props.author = "SoC Architecture / Interconnect / ESL Modeling"
    props.category = "Architecture Review"
    props.comments = "Generated as a design-code-entry architecture review input."
    props.keywords = "SoC, Interconnect, Ring, ESL, OSD, Credit, Backpressure, Performance Modeling"

    add_cover(doc)
    add_executive_summary(doc)
    add_toc(doc)
    add_sections_1_to_4(doc)
    add_sections_5_to_6(doc, names)
    add_section_7(doc, names)
    add_section_8(doc)
    add_section_9(doc, names)
    add_section_10(doc, names)
    add_section_11(doc)
    add_section_12(doc, names)
    add_section_13(doc)
    add_section_14(doc)
    add_section_15(doc)

    for section in doc.sections:
        configure_section(section)
        set_headers_footers(section, cover=False)
    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)
    return DOCX_PATH


def main():
    names = write_diagrams()
    path = build_document(names)
    print(path)


if __name__ == "__main__":
    main()
